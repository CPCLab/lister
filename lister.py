import json
# import elabapy
import os
import pathlib
import platform
import re
import unicodedata
import warnings
from argparse import Namespace
from enum import Enum
from multiprocessing import freeze_support
from pathlib import Path
from pprint import pprint
from typing import Any, Tuple, List, Dict, Union, TypedDict
import PyInstaller
from io import StringIO
import elabapi_python
import latex2mathml.converter
import pandas as pd
import urllib3
import xlsxwriter
from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.shared import Mm, RGBColor
from elabapi_python.rest import ApiException
from gooey import Gooey, GooeyParser
from lxml import etree
from pathvalidate import sanitize_filepath

# TODO: remove the following line when the issue is fixed
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)


# -------------------------------- CLASSES TO HANDLE ENUMERATED CONCEPTS --------------------------------
# Control Flow Metadata Types
class CFMetadata(Enum):
    STEP_TYPE = "step type"
    FLOW_TYPE = "flow type"
    FLOW_PARAMETER = "flow parameter"
    FLOW_LOGICAL_OPERATOR = "flow logical parameter"
    FLOW_COMPARED_VALUE = "flow compared value"
    FLOW_RANGE = "flow range"
    FLOW_OPERATION = "flow operation"
    FLOW_MAGNITUDE = "flow magnitude"
    FLOW_SECTION = "section"
    FLOW_ITERATION_START = "start iteration value"
    FLOW_ITERATION_END = "end iteration value"


class BracketPairErrorMsg(Enum):
    IMPROPER_COMMENT_BRACKET = "ERROR: Mismatch between '(' and ')'. Check line "
    IMPROPER_RANGE_BRACKET = "ERROR: Mismatch between '[' and ']'.  Check line "
    IMPROPER_KEY_VALUE_BRACKET = "ERROR: Mismatch between '{' and '}'.  Check line "
    IMPROPER_FLOW_BRACKET = "ERROR: Mismatch between '<' and '>'.  Check line "


class MiscAlertMsg(Enum):
    ARGUMENT_MISMATCH = "ERROR: Argument type mismatch: numerical value is found while string was expected. " \
                        "Check the value '{0}' in the following set of values: '{1}'."
    UNRECOGNIZED_OPERATOR = "ERROR: The logical operator is not recognized. " \
                            "Please check the operator '{0}' in the following set of values: {1}. " \
                            "Only 'e', 'ne', 'lt', 'lte', 'gt', 'gte' and 'between' are supported."
    UNRECOGNIZED_FLOW_TYPE = "WARNING: The flow type is not recognized. " \
                             "Please check the flow type {0} in the following set of values: {1}."
    RANGE_NOT_TWO_ARGUMENTS = "ERROR: There should only be two numerical arguments on a range separated by a dash " \
                              "(-). Please check the following set of values: {0}."
    RANGE_NOT_NUMBERS = "ERROR: The range values should only contain numbers." \
                        "Check the following part: {0}."
    INVALID_ITERATION_OPERATOR = "ERROR: {0} is not a valid iteration operators. Only +, -, *, / and %% " \
                                 "are supported.Check the following part: {1}."
    IMPROPER_ARGUMENT_NO = "ERROR: Expected number of arguments in the '{0}' statement is {1}, but {2} was found." \
                           "Check the following part: '{3}'"
    ITERATION_OPERATION_NOT_EXIST = "ERROR: The iteration operation is not found, please check the following part: {0}."
    MAGNITUDE_NOT_EXIST = "ERROR: The magnitude of the iteration flow is not found, " \
                          "please check the following part: {0}."
    INACCESSIBLE_RESOURCE = "ERROR: Resource with ID '{0}' is not accessible using the current user's API Token. " \
                            "Please check the resource ID and the user's permission. Reason: {1}, code: {2}, " \
                            "message: {3}, description: {4} Parsing this resource is skipped."
    INACCESSIBLE_EXPERIMENT = "ERROR: Experiment with ID '{0}' is not accessible using the current user's API Token." \
                              " Please check the experiment ID and the user's permission. Reason: {1}, code: {2}, " \
                              "message: {3}, description: {4} Parsing this experiment is skipped."
    SIMILAR_PAR_KEY_FOUND = "WARNING: A combination of similar paragraph number and key has been found, '{0}'. " \
                            "Please make sure that this is intended."
    INACCESSIBLE_ATTACHMENT = "WARNING: File with name '{0}' is not accessible, with the exception: " \
                              "\n {1}. \n Try contacting eLabFTW administrator reporting the exception mentioned."
    INVALID_METADATA_SET_ELEMENT_NO = "ERROR: The number of key value element set must be either two (key-value) " \
                                      "or four (key-value-measure-unit). There are {0} element(s) found in this " \
                                      "key-value set: {1}."
    SINGLE_PAIRED_BRACKET = "WARNING: A Key-Value split with length = 1 is found. This can be caused by a " \
                            "mathematical formula, which is okay and hence no KEY_VALUE pair is written to the " \
                            "metadata. Otherwise please check this pair: {0}."
    MISSING_MML2OMML = "WARNING: Formula is found on the experiment entry. Parsing this formula to docx file " \
                       "requires MML2OMML.XSL file from Microsoft Office to be put on the same directory as " \
                       "config.json file. It is currently downloadable from " \
                       "https://www.exefiles.com/en/xsl/mml2omml-xsl/, Otherwise,  formula parsing is disabled."
    NON_TWO_COLUMNS_LINKED_TABLE = "WARNING: The linked category '{0}' has a table that with {1} column instead of " \
                                   "2. This linked item is skipped. Please recheck and consider using two columns to " \
                                   "allow key-value format."
    NO_HTML_LINE_CONTENT = "WARNING: No HTML line content is found. This can be caused by an empty paragraph. "


class RegexPatterns(Enum):
    EXPLICIT_KEY = r':.+?:'  # catch explicit key which indicated within ":" sign
    SURROUNDED_WITH_COLONS = r'^:.+?:$'  # catch explicit key which indicated within ":" sign
    KEY_VALUE_OR_FLOW = r'\{.+?\}|<.+?>'  # find any occurrences of either KEY_VALUE or control flow
    KEY_VALUE = r'\{.+?\}'  # find any occurrences of KEY_VALUE
    FLOW = r'<.+?>'  # find any occurrences of control flows
    DOI = r"\b(10[.][0-9]{4,}(?:[.][0-9]+)*/(?:(?![\"&\'<>])\S)+)\b"  # catch DOI
    COMMENT = r"\(.+?\)"  # define regex for parsing comment
    FORMULA = r"\$.*\$"  # define regex for parsing formula
    COMMENT_W_CAPTURE_GROUP = r"(\(.+?\)*.*\))"
    COMMENT_VISIBLE = r"\(:(.+?):\)"
    # COMMENT_VISIBLE = "\(:.+?:\)"
    COMMENT_INVISIBLE = r"\(_.+?_\)"
    # catch the end part of KEY_VALUE pairs (the key, tolerating trailing spaces)
    SEPARATOR_AND_KEY = r"\|(\s*\w\s*\.*)+\}"
    BRACKET_MARKUPS = r"([{}<>])"  # catch KEY_VALUE/section bracket annotations
    SEPARATOR_COLON_MARKUP = r"([|:])"  # catch separator and colon annotation
    SEPARATOR_MARKUP = r"([|])"  # catch separator annotation
    PRE_PERIOD_SPACES = r'\s+\.'
    PRE_COMMA_SPACES = r'\s+,'
    SUBSECTION = '(sub)*section'
    SUBSECTION_W_EXTRAS = r'(sub)*section.+'
    # SPAN_ATTR_VAL = r"(\w+-?\w+):(#?\w+?);"
    SPAN_ATTR_VAL = r"(\w+-?\w+):(#?\w+?.?\w+?);"  # catch span attribute value pairs


class ArgNum(Enum):
    ARG_NUM_FOREACH = 2
    ARG_NUM_IF = 4
    ARG_NUM_ELSEIF = 4
    ARG_NUM_ELSE = 1
    ARG_NUM_WHILE = 4
    ARG_NUM_ITERATE = 3
    ARG_NUM_FOR = 5
    ARG_NUM_KV = 2
    ARG_NUM_COMMENT = 1
    ARG_NUM_SECTION = 2


# -------------------------------------------- API Access-related Class -----------------------------------------------
class ApiAccess:

    @classmethod
    def get_resource_item(cls, api_v2_client: elabapi_python.api_client, resource_id: int) -> tuple[
        elabapi_python.Item, str]:
        """
        Get an item from eLabFTW using the resource item ID and API v2 client.

        :param api_v2_client: The API v2 client.
        :param resource_id: The item ID.
        :return: The item (resource) content.
        """
        log = ""
        api_item_response = None
        api_instance = elabapi_python.ItemsApi(api_v2_client)
        print("------------------------------")
        print("Accessing resource item with ID: " + str(resource_id))
        try:
            api_item_response = api_instance.get_item(resource_id, format='json')
        except ApiException as e:
            reason, code, message, description = cls.parse_api_exception(e)
            log = MiscAlertMsg.INACCESSIBLE_RESOURCE.value.format(resource_id, reason, code, message, description)
            print(log)
        return api_item_response, log

    @classmethod
    def parse_api_exception(cls, e: ApiException) -> Tuple[str, str, str, str]:
        """
        Parse an ApiException and return the error details.
        """
        reason = e.reason
        details = e.body.decode('utf-8')  # Decode byte string to string
        details_json = json.loads(details)  # Parse string to JSON
        code = details_json['code']  # Access the code
        message = details_json['message']  # Access the message
        description = details_json['description']  # Access the description
        return reason, code, message, description

    @classmethod
    def get_attachment_long_name(cls, img_path: str) -> str:
        """
        Get an uploads long name from the img path.

        :param img_path: The path of the image.
        :type img_path: str
        :return: The long name of the upload.
        :rtype: str

        This method splits the image path by the separators '&' and '='.
        It then returns the second element of the split path, which corresponds
        to the randomly assigned long name used to access via the URI.
        """
        split_path = GeneralHelper.split_by_separators(img_path, ('&', '='))
        return split_path[1]  # strip first 19 chars to get the long_name field in the upload dictionary

    @classmethod
    def get_exp_title(cls, api_v2_client, exp_item_no: int) -> str:
        """
        Get the title of an experiment from eLabFTW using api_v2_client object the experiment item number.

        :param api_v2_client: eLabFTW API v2 client object
        :param int exp_item_no: eLabFTW experiment item number
        :return: Experiment title as a string
        """
        exp = cls.get_exp(api_v2_client, exp_item_no)
        if exp is None:
            raise ValueError("Failed to retrieve experiment entry.")
        exp_title = exp.__dict__["_title"]
        return exp_title

    @classmethod
    def get_exp_info(cls, experiment: dict) -> List[List[str]]:
        """
        Get experiment information and return it as a list of lists.

        :param experiment: An eLabFTW APIs Experiment object containing experiment information.
        :return: A list of lists containing experiment information in the form of par.no-key-value-measure-units.
        """
        metadata_pairs = [["", "metadata section", "Experiment Info", "", ""],
                          ["", "title", experiment.__dict__["_title"], "", ""],
                          ["", "creation date", experiment.__dict__["_created_at"], "", ""],
                          ["", "category", experiment.__dict__["_type"], "", ""],
                          ["", "author", experiment.__dict__["_fullname"], "", ""],
                          ["", "tags", experiment.__dict__["_tags"], "", ""]]
        return metadata_pairs

    @classmethod
    def get_exp(cls, api_v2_client: elabapi_python.ApiClient, experiment_id: int) -> elabapi_python.Experiment:
        """
        Get an eLab experiment using the provided API client and experiment ID.

        :param elabapi_python.ApiClient api_v2_client: The eLab API client instance.
        :param int experiment_id: The ID of the experiment to retrieve.
        :return: elabapi_python.Experiment exp_response: The retrieved eLab experiment.

        This method uses the provided eLab API client to fetch an experiment with the given ID.
        If an ApiException occurs, it prints the exception message and continues.
        """
        exp_response = None
        api_instance = elabapi_python.ExperimentsApi(api_v2_client)
        print("------------------------------")
        print("Accessing experiment with ID: " + str(experiment_id))
        try:
            exp_response = api_instance.get_experiment(experiment_id, format='json')
        except ApiException as e:
            reason, code, message, description = cls.parse_api_exception(e)
            log = MiscAlertMsg.INACCESSIBLE_EXPERIMENT.value.format(experiment_id, reason, code,
                                                                    message, description)
            print(log)
        return exp_response

    @classmethod
    def get_attachment_ids(cls, exp: Dict, content: Tag) -> Union[list[dict[str, Union[str, Any]]],
                                                            list[Union[str, TypedDict]]]:
        """
        Get upload experiment_id from given experiment and content.
        :param dict exp: a dictionary containing details of an experiment (html body, status, rating, next step, etc.).
        :param bs4.element.Tag content: a bs4 Tag object containing <h1>/<p><img alt=... src=...> Tag that provides the
                link to a particular image file.
        :return: dictionary with keys 'upl_id', 'real_name', and 'hash'
            WHERE
            str upl_id: upload experiment_id of the image attachment, used to access the image through API,
            str real_name: the name of the file when it was uploaded to eLabFTW.
            str hash: the hash of the file when it was uploaded to eLabFTW.
        """
        log = ""
        images = content.find_all('img')
        uploads = exp.__dict__['_uploads']
        results = []
        if len(uploads) > 0:
            try:
                for image in images:
                    upl_long_name = cls.get_attachment_long_name(image['src'])
                    result = {'upl_id': "", 'real_name': "", 'hash': "",
                              "image_path": image['src'], "upl_long_name": upl_long_name}
                    matched_upl = next(upload for upload in uploads if upload.__dict__['_long_name'] == upl_long_name)
                    result['upl_id'] = matched_upl.__dict__['_id']
                    result['real_name'] = matched_upl.__dict__['_real_name']
                    result['hash'] = matched_upl.__dict__['_hash']
                    results.append(result)
            except Exception as e:
                log = MiscAlertMsg.INACCESSIBLE_ATTACHMENT.value.format("NULL", str(e))
                print(log)
                print("Attachment download is skipped...")
                # The dictionary 'result' already has default values set, so no need to set them again here
        return results, log

    @classmethod
    def get_api_v2_endpoint(cls, api_v1_endpoint: str) -> str:
        """
        Convert a version 1 API endpoint to a version 2 API endpoint.

        :param str api_v1_endpoint: version 1 API endpoint.
        :return: str v2endpoint: version 2 API endpoint.
        """
        v2_endpoint = re.sub(r'http://', 'https://', api_v1_endpoint)
        v2_endpoint = re.sub(r'/v1', '/v2', v2_endpoint)
        return v2_endpoint

    @classmethod
    def create_api_v2_client(cls, api_endpoint: str, token: str) -> elabapi_python.ApiClient:
        """
        Create an API v2 client with the given endpoint and token.

        :param api_endpoint: The API endpoint.
        :param token: The API token.
        :return: The API v2 client.
        :rtype: elabapi_python.ApiClient.
        """
        endpoint_v2 = cls.get_api_v2_endpoint(api_endpoint)
        api_v2_config = elabapi_python.Configuration()
        api_v2_config.api_key['api_key'] = token
        api_v2_config.api_key_prefix['api_key'] = 'Authorization'
        api_v2_config.host = endpoint_v2
        api_v2_config.debug = False
        api_v2_config.verify_ssl = False
        api_v2_client = elabapi_python.ApiClient(api_v2_config)
        api_v2_client.set_default_header(header_name='Authorization', header_value=token)
        return api_v2_client

    @classmethod
    def get_save_attachments(cls, path: str, api_v2_client: elabapi_python.ApiClient, exp_id: int) -> str:
        """
        Get a list of attachments in the experiment entry and download these attachments, and return the logs as string.

        :param str path: the path for downloading the attached files, typically named based on experiment title or ID.
        :param elabapi_python.ApiClient api_v2_client: The API v2 client object.
        :param int exp_id: The experiment ID.

        :return log:  The log as a string.
        """

        log = ""

        experiments_api = elabapi_python.ExperimentsApi(api_v2_client)
        uploads_api = elabapi_python.UploadsApi(api_v2_client)
        exp = experiments_api.get_experiment(int(exp_id))

        if platform.system() == "Windows":
            upload_saving_path = path + '\\' + 'attachments'
        else:
            upload_saving_path = path + '/' + 'attachments'

        sanitized_upload_saving_path = sanitize_filepath(upload_saving_path, platform='auto')
        PathHelper.check_and_create_path(sanitized_upload_saving_path)

        # print("LIST OF EXPERIMENT ATTACHMENTS: ")
        # pprint(uploads_api.read_uploads('experiments', exp.id))

        for upload in uploads_api.read_uploads('experiments', exp.id):
            try:
                if platform.system() == "Windows":
                    upload_path = sanitized_upload_saving_path + "\\" + upload.hash + "_" + upload.real_name
                else: # on Unix-based OS
                    upload_path = sanitized_upload_saving_path + "/" + upload.hash + "_" + upload.real_name

                with open(upload_path, 'wb') as file:
                    print("Attachment found: ID: {0}, with name {1}. Writing to {2}.".format(str(upload.id),
                         upload.real_name, upload_path))
                    file.write(uploads_api.read_upload('experiments', exp.id, upload.id, format='binary',
                                                    _preload_content=False).data)
                    file.flush()
            except FileNotFoundError as e:
                print("ERROR: Attachment not found or not accessible: {0}".format(e))
                log = log + "Attachment not found or not accessible: {0}".format(e)

        return log


# ------------------------------------------------ GUI Helper Class ---------------------------------------------------
class GUIHelper:

    @Gooey(optional_cols=0, program_name="LISTER: Life Science Experiment Metadata Parser", default_size=(753, 753),
           navigation="TABBED")
    def parse_gooey_args(self) -> Namespace:
        """
        Get arguments from an existing JSON config to be passed to python Gooey library interface.
        Manual debugging (i.e., printout) is necessary when Gooey is used.

        :returns: args WHERE argparse.Namespace args is an object containing several attributes:
        - str command (e.g., eLabFTW),
        - str output_file_name,
        - int exp_no,
        - str endpoint,
        - str base_output_dir,
        - str token,
        - bool uploadToggle.
        """
        token, endpoint, output_file_name, experiment_no, resource_item_no = self.parse_cfg()
        settings_msg = ('Please ensure to enter the fields below properly, or ask your eLabFTW admin if you '
                        'have questions.')
        parser = GooeyParser(description=settings_msg)
        subs = parser.add_subparsers(help='commands', dest='command')
        base_output_path = PathHelper.get_default_output_path(output_file_name)

        # ELABFTW EXPERIMENT PARAMETERS
        elab_arg_parser = subs.add_parser('parse_experiment', prog="Parse Experiment",
                                          help='Parse metadata from an eLabFTW experiment entry')
        io_args = elab_arg_parser.add_argument_group("Input/Output Arguments", gooey_options={'columns': 1})
        radio_group = io_args.add_mutually_exclusive_group(required=True,
                                                           gooey_options={'title': "Naming method for the outputs",
                                                                          'initial_selection': 0})
        radio_group.add_argument("-t", "--title", metavar="Title", action="store_true",
                                 help='Name files and folders based on experiment title')
        radio_group.add_argument("-i", "--experiment_id", metavar="ID", action="store_true",
                                 help='Name files and folders based on the experiment ID')

        io_args.add_argument('base_output_dir', metavar='Base output directory',
                             help='Local directory generally used to save your outputs', type=str,
                             default=base_output_path, widget='DirChooser')
        elabftw_args = elab_arg_parser.add_argument_group("eLabFTW Arguments", gooey_options={'columns': 2})
        elabftw_args.add_argument('exp_no', metavar='eLabFTW experiment ID',
                                  help='Integer indicated in the URL of the experiment',
                                  default=experiment_no, type=int)
        elabftw_args.add_argument('endpoint', metavar="eLabFTW API endpoint URL",
                                  help='Ask your eLabFTW admin to provide the endpoint URL for you', default=endpoint,
                                  type=str)
        elabftw_args.add_argument('token', metavar='eLabFTW API Token',
                                  help='Ask your eLabFTW admin to generate an API token for you', default=token,
                                  type=str)

        # ELABFTW resource PARAMETERS
        elab_arg_parser = subs.add_parser('parse_resource', prog="Parse Container",
                                          help='Parse metadata from an eLabFTW resource/container items')
        io_args = elab_arg_parser.add_argument_group("Input/Output Arguments", gooey_options={'columns': 1})
        radio_group = io_args.add_mutually_exclusive_group(required=True,
                                                           gooey_options={'title': "Naming method for the outputs",
                                                                          'initial_selection': 0})
        radio_group.add_argument("-t", "--title", metavar="Title", action="store_true",
                                 help='Name files and folders based on container type + title, including the '
                                      'underlying experiments')
        radio_group.add_argument("-i", "--experiment_id", metavar="ID", action="store_true",
                                 help='Name files and folders based on container type + ID, including the '
                                      'underlying experiments')

        io_args.add_argument('base_output_dir', metavar='Base output directory',
                             help='Local directory generally used to save your outputs', type=str,
                             default=base_output_path, widget='DirChooser')
        elabftw_args = elab_arg_parser.add_argument_group("eLabFTW Arguments", gooey_options={'columns': 2})
        elabftw_args.add_argument('resource_item_no', metavar='eLabFTW Resource/Container Item ID',
                                  help='Integer indicated in the URL of the resource/container item',
                                  default=resource_item_no, type=int)
        elabftw_args.add_argument('endpoint', metavar="eLabFTW API endpoint URL",
                                  help='Ask your eLabFTW admin to provide the endpoint URL for you', default=endpoint,
                                  type=str)
        elabftw_args.add_argument('token', metavar='eLabFTW API Token',
                                  help='Ask your eLabFTW admin to generate an API token for you', default=token,
                                  type=str)

        args = parser.parse_args()
        return args

    @classmethod
    def parse_cfg(cls) -> Tuple[str, str, str, int, int]:
        """
        Parse JSON config file, requires existing config.json file which should be specified on certain directory.

        The directory is OS-dependent. On Windows/Linux, it is in the same folder as the script/executables.
        On macOS, it is in the users' Apps/lister/config.json file.

        :returns: tuple (token, endpoint, output_file_name, experiment_id)
            str token: eLabFTW API Token,
            str endpoint: eLabFTW API endpoint URL,
            str output_file_name: filename to be used for all the outputs (xlsx/json metadata, docx documentation,
                                  log file),
            int experiment_id: the parsed experiment ID (int).
            int resource_item_no: the parsed resource/container item ID (int).
        """

        input_file = PathHelper.manage_input_path() + "config.json"
        print("CONFIG FILE: %s" % input_file)
        # using ...with open... allows file to be closed automatically.
        with open(input_file, encoding="utf-8") as json_data_file:
            data = json.load(json_data_file)
        token = data['elabftw']['token']
        endpoint = data['elabftw']['endpoint']
        experiment_id = data['elabftw']['exp_no']
        output_file_name = data['elabftw']['output_file_name']
        resource_item_no = data['elabftw']['resource_item_no']
        return token, endpoint, output_file_name, experiment_id, resource_item_no


# -------------------------------------------- File serialization Class ------------------------------------------------

class Serializer:

    @classmethod
    def write_to_docx(cls, exp: dict, path: str) -> str:
        """
        fetch an experiment, clean the content from LISTER annotation markup and serialize the result to a docx file.

        :param dict exp: dictionary containing the properties of the experiment, including its HTML body content.
        :param str path: the path for writing the docx file, typically named based on experiment title or ID.

        :return: str log: log of the process.
        """
        document = Document()
        all_references = []
        log = ""
        tagged_contents = TextCleaner.get_nonempty_body_tags(exp)

        watched_tags = ['div', 'p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'span', 'strong', 'sub', 'em', 'sup']
        for content in tagged_contents:  # iterate over a list of tags

            if isinstance(content, Tag):

                if content.name == "div":
                    print("A div tag is found, unwrapping...")
                    # print("The content is a div, unwrapping...")
                    content.unwrap()

                if len(content.select("img")) > 0:
                    # upl_id, real_name, hash = ApiAccess.get_attachment_ids(experiment, content)
                    image_ids, log = ApiAccess.get_attachment_ids(exp, content)
                    for image_id in image_ids:
                        # print(str(image_id['real_name']), path)
                        DocxHelper.add_img_to_doc(document, image_id['real_name'], path, image_id['hash'])
                elif any(x in content.name for x in watched_tags):
                    references, log = DocxHelper.write_tag_to_doc(document, content)
                    if len(references) > 0:
                        all_references.extend(references)
                if content.name == "table":
                    print("A table is found, writing to docx...")
                    DocxHelper.add_table_to_doc(document, content)
                if content.name == "img":
                    image_ids, log = ApiAccess.get_attachment_ids(exp, content)
                    for image_id in image_ids:
                        # print(str(image_id['real_name']), path)
                        DocxHelper.add_img_to_doc(document, image_id['real_name'], path, image_id['hash'])
        if len(all_references) > 0:
            document.add_heading("Reference", level=1)
            for reference in all_references:
                document.add_paragraph(reference, style='List Number')
        try:
            if platform.system() == "Windows":
                docx_path = path + "\\" + PathHelper.derive_filename_from_experiment(exp) + '.docx'
                docx_path = sanitize_filepath(docx_path, platform="auto")
            else:
                docx_path = path + '/' + PathHelper.derive_filename_from_experiment(exp) + '.docx'
            document.save(docx_path)
        except FileNotFoundError as e:
            print("ERROR: DOCX file cannot be written: {0}".format(e))
            log = log + "ERROR: DOCX file cannot be written: {0}".format(e)
        return log

    # Used to serialize extracted metadata to json file.
    @classmethod
    def write_to_json(cls, lst: List, exp: dict, path: str) -> None:
        """
        Write a list to a JSON file.
        :param lst: The list to write to the JSON file.
        :param exp: The experiment title or ID.
        :param path: The path for writing the JSON file.
        """
        filename = f"{PathHelper.derive_filename_from_experiment(exp)}.json"
        try:
            if platform.system() == "Windows":
                json_path = path + "\\" + filename
                json_path = sanitize_filepath(json_path, platform="auto")
            else:
                json_path = path + '/' + filename
            with open(json_path, "w", encoding="utf-8") as f:
                json.dump(lst, f, ensure_ascii=False)
        except FileNotFoundError as e:
            print("ERROR: JSON file cannot be written: {0}".format(e))

    # Used to write into the log file.
    # def write_log(log, full_path=output_path_and_filename):
    @classmethod
    def write_log(cls, log_text: str, path: str) -> None:
        """
        Write the log to a file.
        :param log_text: The log to be written to the file.
        :param path: The path for writing the log file.
        """
        log_text = log_text.strip()
        PathHelper.check_and_create_path(path)
        try:
            if platform.system() == "Windows":
                log_path = path + "\\" + 'lister-report.log'
                log_path = sanitize_filepath(log_path, platform="auto")
            else:
                log_path = path + '/' + 'lister-report.log'
            with open(log_path, "w", encoding="utf-8") as f:
                f.write(log_text)
        except FileNotFoundError as e:
            print("ERROR: LOG file cannot be written: {0}".format(e))

    @classmethod
    def write_to_xlsx(cls, metadata_set: List, exp: dict, path: str) -> None:
        """
        Write extracted order/key/value/measure/unit to an Excel file.

        :param list metadata_set: list containing the order (paragraph number)/key/value/measure/unit to be written.
        :param dict exp: an experiment object.
        :param str path: the path for writing the xlsx file, typically named based on experiment title or ID.
        """
        PathHelper.check_and_create_path(path)
        header = ["PARAGRAPH NUMBER", "KEY", "VALUE", "MEASURE", "UNIT"]
        # json.dump(list, open(path + '/' + derive_filename_from_exp(experiment) + ".json", 'w', encoding="utf-8"),
        # ensure_ascii=False)
        try:
            if platform.system() == "Windows":
                xlsx_path = path + "\\" + PathHelper.derive_filename_from_experiment(exp) + ".xlsx"
                xlsx_path = sanitize_filepath(xlsx_path, platform="auto")
            else:
                xlsx_path = path + '/' + PathHelper.derive_filename_from_experiment(exp) + ".xlsx"
            with xlsxwriter.Workbook(xlsx_path) as workbook:
                # formatting cells
                header_format = workbook.add_format({'bold': True, 'bg_color': '9bbb59', 'font_color': 'ffffff'})
                default_format = workbook.add_format({'border': 1, 'border_color': '9bbb59'})
                section_format = workbook.add_format({'border': 1, 'border_color': '9bbb59', 'bg_color': 'ebf1de'})
                # creating and formatting worksheet
                worksheet = workbook.add_worksheet()
                worksheet.write_row(0, 0, header, header_format)
                worksheet.set_column('A:A', 19)
                worksheet.set_column('B:B', 18)
                worksheet.set_column('C:C', 30)
                worksheet.set_column('D:E', 15)
                for row_no, data in enumerate(metadata_set):
                    key = data[1]
                    # do not use regex here, or it will be very slow
                    # if re.match(RegexPatterns.SUBSECTION.value, data[1].lower()):
                    if len(key) >= 7 and key[0:7].casefold() == "section".casefold() or key.casefold() == "metadata section":
                        worksheet.write_row(row_no + 1, 0, data, section_format)
                    else:
                        try:
                            worksheet.write_row(row_no + 1, 0, data, default_format)
                        except xlsxwriter.exceptions.FileCreateError as e:
                            print(e)
                           #  print("ERROR: Excel file cannot be written: {0}".format(e))
        except xlsxwriter.exceptions.FileCreateError as e:
            print(e)





# ---------------------------------------------- Metadata Extraction Class --------------------------------------------


class MetadataExtractor:

    @classmethod
    def is_explicit_key(cls, key: str) -> bool:
        """
        Check whether the string is an explicit key.

        :param str key: checked string.
        :return: bool stating whether the key is a LISTER explicit key.

        """
        if re.match(RegexPatterns.EXPLICIT_KEY.value, key):
            return True
        else:
            return False

    @classmethod
    def extract_flow_type(cls, paragraph_no: int, flow_control_pair: str) -> Tuple[List[List], str, bool]:
        """
        Extracts the type of flow found on any annotation with angle brackets, which can be control flow or sectioning.

        :param int paragraph_no: paragraph number on where the control flow fragment string was found.
        :param str flow_control_pair: the control-flow pair string to be extracted for metadata.
        :returns: tuple (key_value, flow_log, is_error)
            WHERE
            list key_value: list of list, each list contains a full complete control flow metadata line
                        e.g. [['-', 'section level 0', 'Precultures', '', '']],
            str flow_log: log resulted from running this and subsequent functions,
            bool is_error: flag that indicates whether an error occurred.
        """
        flow_log = ""
        # print("flow_control_pair: " + str(flow_control_pair))
        key_value = []
        cf = flow_control_pair[1:-1]
        control_flow_split = re.split(r"\|", cf)
        flow_type = control_flow_split[0]
        flow_type = flow_type.strip()
        flow_type = flow_type.lower()
        if flow_type == "for each":
            key_value, log, is_error = cls.process_foreach(paragraph_no, control_flow_split)
            if log != "":
                flow_log = flow_log + "\n" + log
        elif flow_type == "while":
            key_value, log, is_error = cls.process_while(paragraph_no, control_flow_split)
            if log != "":
                flow_log = flow_log + "\n" + log
        elif flow_type == "if":
            key_value, log, is_error = cls.process_if(paragraph_no, control_flow_split)
            if log != "":
                flow_log = flow_log + "\n" + log
        elif flow_type == "else if" or flow_type == "elif":
            key_value, log, is_error = cls.process_elseif(paragraph_no, control_flow_split)
            if log != "":
                flow_log = flow_log + "\n" + log
        elif flow_type == "else":
            key_value, log, is_error = cls.process_else(paragraph_no, control_flow_split)
            if log != "":
                flow_log = flow_log + "\n" + log
        elif flow_type == "for":
            key_value, log, is_error = cls.process_for(paragraph_no, control_flow_split)
            if log != "":
                flow_log = flow_log + "\n" + log
        # elif flow_type.casefold() == "section".casefold():
        elif re.match(RegexPatterns.SUBSECTION.value, flow_type, flags=re.IGNORECASE):
            key_value, log, is_error = cls.process_section(control_flow_split)
            if log != "":
                flow_log = flow_log + "\n" + log
        elif flow_type == "iterate":
            key_value, log, is_error = cls.process_iterate(paragraph_no, control_flow_split)
            if log != "":
                flow_log = flow_log + "\n" + log
        else:
            is_error = True
            log = MiscAlertMsg.UNRECOGNIZED_FLOW_TYPE.value.format(control_flow_split[0].upper(),
                                                                   control_flow_split)  # + "\n"
            print(log)
            flow_log = flow_log + "\n" + log
            # print(flow_log)
        # print("key_value: " + str(key_value) + "\n\n")
        return key_value, flow_log, is_error

    @classmethod
    def process_section(cls, control_flow_split: List[str]) -> Tuple[List[List], str, bool]:
        """
        Convert key value based on a section to a full section metadata entry

        :param list control_flow_split: list of strings split e.g., ['Section', 'Remarks']
        :returns: tuple (key_value, log, is_error)
            WHERE
            list key_value: list of list, each list contains a full section-metadata line
                        e.g. [['-', 'section level 0', 'Precultures', '', '']],
            str log: log resulted from running this and subsequent functions,
            bool is_error: flag that indicates whether an error occurred.
        """
        key_value = []
        section_log = ""
        is_error = False
        log, is_sect_error = Validator.validate_section(control_flow_split)
        if is_sect_error:
            is_error = True
            section_log = section_log + "\n" + log
        else:
            section_keyword = control_flow_split[0].lower()
            section_level = section_keyword.count("sub")
            key_value.append(
                ["-", CFMetadata.FLOW_SECTION.value + " level " + str(section_level), control_flow_split[1], '', ''])
        return key_value, section_log, is_error

    @classmethod
    def process_ref_resource_item(cls, api_v2_client: elabapi_python.ApiClient, item_api_response) -> None:
        """
        Process reference resource item, using the initial resource ID for that container item (e.g., publication).

        :param api_v2_client: An instance of the API v2 client object, containing eLabFTW API-related information.
        :param item_api_response: The API response of the reference resource item.
        :return: None
        """

        # TODO: also get the list of related experiments instead of linked experiments only,
        #  status: pending. see https://github.com/elabftw/elabftw/issues/4811
        try:
            experiments = item_api_response.__dict__["_experiments_links"]
            for experiment in experiments:
                exp_path = output_path + PathHelper.slugify(experiment.__dict__["_title"])
                cls.process_experiment(api_v2_client, experiment.__dict__["_itemid"], exp_path)
        except ApiException as e:
            print("Exception when calling ItemsApi->getItem: %s\n" % e)

    @classmethod
    def process_linked_resource_item_api_v2(cls, api_v2_client: elabapi_python.ApiClient, resource_id: int) -> \
            Tuple[Union[List[List[str]], str], str]:
        """
        Process a linked resource item and return its metadata and log.

        :param elabapi_python.ApiClient api_v2_client: An instance of the API v2 client object, containing eLabFTW
        API-related information.
        :param resource_id: The ID of the linked resource item.
        :return: A tuple containing the resource item metadata and log.
        """
        api_instance = elabapi_python.ItemsApi(api_v2_client)

        try:
            # Read an item
            linked_item = api_instance.get_item(resource_id)
            html_body = getattr(linked_item, 'body')
            # category = getattr(linked_item, 'mainattr_title') # only works for elabapi-python 0.4.1.
            category = getattr(linked_item, 'category_title')

            dfs = pd.read_html(StringIO(html_body))
            df = pd.concat(dfs)
            df_col_no = df.shape[1]
            log = ""
            if df_col_no != 2:
                log = MiscAlertMsg.NON_TWO_COLUMNS_LINKED_TABLE.value.format(category, df_col_no) + "\n"
                print(log)
                resource_item_metadata_set = None
            else:
                df.columns = ["metadata section", category]
                df.insert(loc=0, column="", value="")
                df = df.reindex(df.columns.tolist() + ['', ''], axis=1)  # add two empty columns
                filtered_df = df.fillna('')  # fill empty cells with empty string
                resource_item_metadata_set = [filtered_df.columns.values.tolist()] + filtered_df.values.tolist()
        except ApiException as e:
            resource_item_metadata_set = ""
            log = "Exception when calling ItemsApi->getItem: %s\n" % e
            print(log)
        return resource_item_metadata_set, log

    @classmethod
    def process_experiment(cls, api_v2_client: elabapi_python.ApiClient, exp_id: int, path: str) -> None:
        """
        Process an experiment and save its information to various formats.

        :param elabapi_python.ApiClient api_v2_client: The API v2 client.
        :param int exp_id: The experiment number.
        :param str path: The path for saving the output files.
        """
        overall_log = ""

        experiment_instance = elabapi_python.ExperimentsApi(api_v2_client)
        # experiment_response = experiment_instance.get_experiment(int(exp_no))

        print("------------------------------")
        print("Accessing experiment with ID: " + str(exp_id))
        try:
            experiment_response = experiment_instance.get_experiment(int(exp_id))
            linked_resources = experiment_response.__dict__['_items_links']
            # get the IDs of the linked resources
            linked_resource_ids = [linked_resource.__dict__["_itemid"] for linked_resource in linked_resources]

            # get the respective category of the linked resources
            id_and_category = {}
            excluded_item_types = ["MM", "Publication", "Protocols", "Protocol", "Methods", "Method", "Recipe"]

            # this will only work with elabapi-python 0.4.1.
            # unfortunately, the response from the API is not consistent between versions, so it may be a good idea to
            # fix the version of elabapi-python to a specific version in the requirements.txt in the future.
            # for linked_resource in linked_resources:
            # id_and_category[linked_resource.__dict__["_itemid"]] = linked_resource.__dict__["_mainattr_title"]

            print("---------------- linked_resource_ids: ---------------- ")
            pprint(linked_resource_ids)

            for linked_resource_id in linked_resource_ids:
                # get the linked resource item by ID
                linked_resource, resource_log = ApiAccess.get_resource_item(api_v2_client, linked_resource_id)
                overall_log = overall_log + "\n" + resource_log
                # pprint(linked_resource)
                if linked_resource is not None:
                    id_and_category[linked_resource.__dict__["_id"]] = linked_resource.__dict__["_category_title"]
            # pprint(id_and_category)

            filtered_id_and_category = {key: value for key, value in id_and_category.items() if value.lower() not in
                                        [item.lower() for item in excluded_item_types]}
            # pprint(filtered_id_and_category)

            overall_metadata_set = []
            # the 'key' here is the ID of the resource item.
            for key in filtered_id_and_category:
                resource_item_metadata_set, log = MetadataExtractor.process_linked_resource_item_api_v2(api_v2_client,
                                                                                                        key)
                overall_log = overall_log + "\n" + log
                overall_metadata_set.extend(resource_item_metadata_set)

            experiment_metadata_info = ApiAccess.get_exp_info(experiment_response)
            overall_metadata_set.extend(experiment_metadata_info)
            experiment_metadata_set, log = (MetadataExtractor.conv_html_to_metadata
                                            (experiment_response.__dict__["_body"]))
            overall_log = overall_log + "\n" + log
            overall_metadata_set.extend(experiment_metadata_set)

            log = ApiAccess.get_save_attachments(path, api_v2_client, int(exp_id))
            overall_log = overall_log + "\n" + log
            docx_log = Serializer.write_to_docx(experiment_response, path)
            overall_log = overall_log + "\n" + docx_log

            Serializer.write_to_json(overall_metadata_set, experiment_response, path)
            Serializer.write_to_xlsx(overall_metadata_set, experiment_response, path)
            Serializer.write_log(overall_log, path)

        except ApiException as e:
            reason, code, message, description = ApiAccess.parse_api_exception(e)
            exp_log = MiscAlertMsg.INACCESSIBLE_EXPERIMENT.value.format(id, reason, code, message, description)
            print(exp_log)
            Serializer.write_log(exp_log, path)

    # only process the comment within (key value measure unit) pairs and remove its content
    # (unless if it is begun with "!")
    @classmethod
    def process_internal_comment(cls, string_with_brackets: str) -> Tuple[str, str]:
        """
        Separates actual part of a lister bracket annotation fragment (key/value/measure/unit) with the
        trailing comments.

        Internal comment refers to any comment that is available within a fragment of a lister bracket annotation.
        Internal comment will not be bypassed to the metadata output.
        However, internal comment is important to be provided to make the experiment clear-text readable in the
        docx output.

        :param str string_with_brackets: a lister bracket annotation fragment with a comment.
        :returns: tuple (actual_fragment, internal_comment)
            WHERE
            str actual_fragment:  string containing the actual element of metadata, it can be either
                                  key/value/measure/unit,
            str internal_comment: string containing the comment part of the string fragment, with brackets retained.
        """
        comment = re.search(RegexPatterns.COMMENT.value, string_with_brackets)
        comment = comment.group(0)
        remains = string_with_brackets.replace(comment, '')
        actual_fragment, internal_comment = remains.strip(), comment.strip()
        return actual_fragment, internal_comment

    @classmethod
    def process_foreach(cls, paragraph_no: int, control_flow_split: List[str]) -> Tuple[List[List], str, bool]:
        """
        Converts key-value based on foreach control-metadata entry.

        :param int paragraph_no: paragraph number where string fragment containing the referred pair was found.
        :param list control_flow_split: list of split string.
        :returns: tuple (key_value, log, is_error)
            WHERE
            list key_value: list of list, each list contains a full control-flow metadata,
            str log: log resulted from running this and subsequent functions,
            bool is_error: flag that indicates whether an error occurred.
        """
        key_value = []
        log, is_error = Validator.validate_foreach(control_flow_split)
        if is_error:
            print(log)
        step_type = "iteration"
        key_value.append([paragraph_no, CFMetadata.STEP_TYPE.value, step_type, '', ''])
        flow_type = control_flow_split[0]
        key_value.append([paragraph_no, CFMetadata.FLOW_TYPE.value, flow_type, '', ''])
        flow_param = control_flow_split[1]
        key_value.append([paragraph_no, CFMetadata.FLOW_PARAMETER.value, flow_param, '', ''])
        return key_value, log, is_error

    @classmethod
    def process_while(cls, paragraph_no: int, control_flow_split: List[str]) -> Tuple[List[List], str, bool]:
        """
        Convert key value based on while control-metadata entry.

        :param int paragraph_no: paragraph number where string fragment containing the referred pair was found.
        :param list control_flow_split: list of split string.
        :returns: tuple (key_value, log, is_error)
            WHERE
            list key_value: list of list, each list contains a full control-flow metadata,
            str log: log resulted from running this and subsequent functions,
            bool is_error: flag that indicates whether an error occurred.
        """
        key_value = []
        log, is_error = Validator.validate_while(control_flow_split)
        if is_error:
            print(log)
        try:
            step_type = "iteration"
            key_value.append([paragraph_no, CFMetadata.STEP_TYPE.value, step_type, '', ''])
            flow_type = control_flow_split[0]
            key_value.append([paragraph_no, CFMetadata.FLOW_TYPE.value, flow_type, '', ''])
            flow_param = control_flow_split[1]
            key_value.append([paragraph_no, CFMetadata.FLOW_PARAMETER.value, flow_param, '', ''])
            flow_logical_operator = control_flow_split[2]
            key_value.append([paragraph_no, CFMetadata.FLOW_LOGICAL_OPERATOR.value, flow_logical_operator, '', ''])
            flow_compared_value = control_flow_split[3]
            key_value.append([paragraph_no, CFMetadata.FLOW_COMPARED_VALUE.value, flow_compared_value, '', ''])
        except IndexError as e:
            is_error = True
            print(f"Iteration error occurred: {e}")
        return key_value, log, is_error

    @classmethod
    def process_if(cls, paragraph_no: int, control_flow_split: List[str]) -> Tuple[List[List], str, bool]:
        """
        Convert key-value based on if control-metadata entry.

        :param int paragraph_no: paragraph number where string fragment containing the referred pair was found.
        :param list control_flow_split: list of split string.
        :returns: tuple (key_value, log, is_error)
            WHERE
            list key_value: list of list, each list contains a full control-flow metadata,
            str log: log resulted from running this and subsequent functions,
            bool is_error: flag that indicates whether an error occurred.
        """
        key_value = []
        log, is_error = Validator.validate_if(control_flow_split)
        if is_error:
            print(log)
        step_type = "conditional"
        key_value.append([paragraph_no, CFMetadata.STEP_TYPE.value, step_type, '', ''])
        flow_type = control_flow_split[0]
        key_value.append([paragraph_no, CFMetadata.FLOW_TYPE.value, flow_type, '', ''])
        flow_param = control_flow_split[1]
        key_value.append([paragraph_no, CFMetadata.FLOW_PARAMETER.value, flow_param, '', ''])
        flow_logical_operator = control_flow_split[2]
        key_value.append([paragraph_no, CFMetadata.FLOW_LOGICAL_OPERATOR.value, flow_logical_operator, '', ''])
        flow_compared_value = control_flow_split[3]
        key_value.append([paragraph_no, CFMetadata.FLOW_COMPARED_VALUE.value, flow_compared_value, '', ''])
        return key_value, log, is_error

    @classmethod
    def process_elseif(cls, paragraph_no: int, control_flow_split: List[str]) -> Tuple[List[List], str, bool]:
        """
        Convert key-value based on else-if control-metadata entry.

        :param int paragraph_no: paragraph number where string fragment containing the referred pair was found.
        :param list control_flow_split: list of split string.
        :returns: tuple (key_value, log, is_error)
            WHERE
            list key_value: list of list, each list contains a full control-flow metadata,
            str log: log resulted from running this and subsequent functions,
            bool is_error: flag that indicates whether an error occurs.
        """
        key_value = []
        log, is_error = Validator.validate_elseif(control_flow_split)
        if is_error:
            print(log)
        step_type = "conditional"
        key_value.append([paragraph_no, CFMetadata.STEP_TYPE.value, step_type, '', ''])
        flow_type = control_flow_split[0]
        key_value.append([paragraph_no, CFMetadata.FLOW_TYPE.value, flow_type, '', ''])
        flow_param = control_flow_split[1]
        key_value.append([paragraph_no, CFMetadata.FLOW_PARAMETER.value, flow_param, '', ''])
        flow_logical_operator = control_flow_split[2]
        key_value.append([paragraph_no, CFMetadata.FLOW_LOGICAL_OPERATOR.value, flow_logical_operator, '', ''])
        flow_compared_value = control_flow_split[3]
        if re.search(r"\[.*?\]", flow_compared_value):
            key_value.append([paragraph_no, CFMetadata.FLOW_RANGE.value, flow_compared_value, '', ''])
            start, end, range_log, range_is_error = cls.process_range(flow_compared_value)
            key_value.append([paragraph_no, CFMetadata.FLOW_ITERATION_START.value, start, '', ''])
            key_value.append([paragraph_no, CFMetadata.FLOW_ITERATION_END.value, end, '', ''])
        else:
            key_value.append([paragraph_no, CFMetadata.FLOW_COMPARED_VALUE.value, flow_compared_value, '', ''])
        return key_value, log, is_error

    @classmethod
    def process_else(cls, paragraph_no: int, control_flow_split: List[str]) -> Tuple[List[List], str, bool]:
        """
        Convert key value based on else control-metadata entry.

        :param int paragraph_no: paragraph number where string fragment containing the referred pair was found.
        :param list control_flow_split: list of split string.
        :returns: tuple (key_value, log, is_error)
            WHERE
            list key_value: list of list, each list contains a full control-flow metadata,
            str log: log resulted from running this and subsequent functions,
            bool is_error: flag that indicates whether an error occurred.
        """
        print(control_flow_split)
        key_value = []
        log, is_error = Validator.validate_else(control_flow_split)
        if is_error:
            print(log)
        step_type = "conditional"
        key_value.append([paragraph_no, CFMetadata.STEP_TYPE.value, step_type, '', ''])
        flow_type = control_flow_split[0]
        key_value.append([paragraph_no, CFMetadata.FLOW_TYPE.value, flow_type, '', ''])
        return key_value, log, is_error

    @classmethod
    def process_range(cls, flow_range: str) -> Tuple[float, float, str, bool]:
        """
        Processes a given flow range string and returns the range values, log, and error status. This function
        validates the flow range string using a Validator, logs any errors encountered during validation, and then
        extracts the numerical range values from the validated string. If the validation fails, it prints the error
        log. Otherwise, it splits the flow range string to extract the start and end values of the range.


        :param str flow_range: A string representing the flow range in the format "[start-end]", where "start" and
                               "end" are numerical values.
        :returns: Tuple[float, float, str, bool]: A tuple containing the start and end values of the range as floats,
         the log message as a string, and a boolean indicating whether an error occurred during processing.

            WHERE
            list range_values: list of list, each list contains a full control-flow metadata,
            str log: log resulted from running this and subsequent functions,
            bool is_error: flag that indicates whether an error occurred.
        """
        log, is_error = Validator.validate_range(flow_range)
        if is_error:
            print(log)
            range_values = None
        else:
            range_values = re.split("-", flow_range[1:-1])
        return float(range_values[0]), float(range_values[1]), log, is_error

    @classmethod
    def process_for(cls, paragraph_no: int, control_flow_split: List[str]) -> Tuple[List[List], str, bool]:
        """
        Convert key value based on for control-metadata entry. Please consult LISTER documentation on GitHub.

        :param int paragraph_no: paragraph number where string fragment containing the referred pair was found.
        :param list control_flow_split: list of split string.
        :returns: tuple (key_value, log, is_error)
            WHERE
            list key_value: list of list, each list contains a full control-flow metadata,
            str log: log resulted from running this and subsequent functions,
            bool is_error: flag that indicates whether an error occurred.
        """
        key_value = []
        for_log = ""
        is_error = False
        for_validation_log, is_for_error = Validator.validate_for(control_flow_split)
        if is_for_error:
            for_log = for_log + "\n" + for_validation_log
            is_error = True
            print(for_validation_log)
        step_type = "iteration"
        key_value.append([paragraph_no, CFMetadata.STEP_TYPE.value, step_type, '', ''])
        flow_type = control_flow_split[0]
        key_value.append([paragraph_no, CFMetadata.FLOW_TYPE.value, flow_type, '', ''])
        flow_param = control_flow_split[1]
        key_value.append([paragraph_no, CFMetadata.FLOW_PARAMETER.value, flow_param, '', ''])
        flow_range = control_flow_split[2]
        key_value.append([paragraph_no, CFMetadata.FLOW_RANGE.value, flow_range, '', ''])
        start, end, range_log, is_range_error = cls.process_range(flow_range)
        if is_range_error:
            for_log = for_log + "\n" + range_log
            print(range_log)
            is_error = True
        key_value.append([paragraph_no, CFMetadata.FLOW_ITERATION_START.value, start, '', ''])
        key_value.append([paragraph_no, CFMetadata.FLOW_ITERATION_END.value, end, '', ''])
        try:
            flow_operation = control_flow_split[3]
            key_value.append([paragraph_no, CFMetadata.FLOW_OPERATION.value, flow_operation, '', ''])
        except KeyError as e:
            is_error = True
            print(f"Iteration error occurred: {e}")
            print(MiscAlertMsg.ITERATION_OPERATION_NOT_EXIST.value.format(control_flow_split))
            for_log = for_log + "\n" + MiscAlertMsg.ITERATION_OPERATION_NOT_EXIST.value.format(control_flow_split)
        try:
            print(control_flow_split)
            print(type(control_flow_split))
            print(len(control_flow_split))
            # print(len(control_flow_split))
            flow_magnitude = control_flow_split[4]
            print("flow_magnitude: ")
            pprint(flow_magnitude)
            key_value.append([paragraph_no, CFMetadata.FLOW_MAGNITUDE.value, flow_magnitude, '', ''])
        except IndexError as e:
            is_error = True
            print(f"Iteration error occurred: {e}")
            print(MiscAlertMsg.MAGNITUDE_NOT_EXIST.value.format(control_flow_split))
            for_log = for_log + "\n" + MiscAlertMsg.MAGNITUDE_NOT_EXIST.value.format(control_flow_split)
        return key_value, for_log, is_error

    # should happen only after having 'while' iterations to provide additional steps on the iterator
    @classmethod
    def process_iterate(cls, paragraph_no: int, control_flow_split: List[str]) -> Tuple[List[List], str, bool]:
        """
        Convert key value based on while control-metadata entry. Please consult LISTER documentation on GitHub.

        :param int paragraph_no: paragraph number where string fragment containing the referred pair was found.
        :param list control_flow_split: list of split string.
        :returns: tuple (key_value, log, is_error)
            WHERE
            list key_value: list of list, each list contains a full control-flow metadata,
            str log: log resulted from running this and subsequent functions,
            bool is_error: flag that indicates whether an error occurred.
        """
        key_value = []
        iterate_log = ""
        log, is_error = Validator.validate_iterate(control_flow_split)
        if is_error:
            iterate_log = iterate_log + "\n" + log
            print(iterate_log)
        flow_type = control_flow_split[0]
        key_value.append([paragraph_no, CFMetadata.FLOW_TYPE.value, flow_type + "  (after while)"])
        flow_operation = control_flow_split[1]
        key_value.append([paragraph_no, CFMetadata.FLOW_OPERATION.value, flow_operation])
        try:
            flow_magnitude = control_flow_split[2]
            key_value.append([paragraph_no, CFMetadata.FLOW_MAGNITUDE.value, flow_magnitude])
        except IndexError as e:
            is_error = True
            print(f"Iteration error occurred: {e}")
            print(MiscAlertMsg.MAGNITUDE_NOT_EXIST.value.format(control_flow_split))
            iterate_log = iterate_log + "\n" + MiscAlertMsg.MAGNITUDE_NOT_EXIST.value.format(control_flow_split)
        return key_value, iterate_log, is_error

    @classmethod
    def parse_lines_to_metadata(cls, lines: List[str]) -> Tuple[List, List[str], str]:
        """
        Get a list of metadata pairs [order, key, value, measure, unit] or ['-', sec. level, section name, '', '']
        from nbsp-clean lines.
        :param list lines: list of lines cleaned up from nbsp.
        :return: tuple (metadata_pairs, internal_comments, log)
            WHERE
            list metadata_pairs: list of [order, key, value, measure, unit] or
                                ['-', section level, section name, '', ''],
            str internal_comments: placeholder for found internal comments within key-value pairs - currently unused,
            str log: log from running subsequent functions.
        """
        paragraph_no = 0
        metadata_pairs = []
        metadata_header = ["", "metadata section", "Experiment Context", "", ""]
        metadata_pairs.append(metadata_header)
        paragraph_key_pairs = []
        log = ""
        internal_comments = []
        for line in lines:
            # Check bracketing validity
            bracketing_log, is_bracket_error = Validator.check_bracket_num(paragraph_no, line)
            log = log + bracketing_log  # + "\n"
            if is_bracket_error:
                break

            # Extract KEY_VALUE and flow metadata
            key_value_and_flow_pairs = re.findall(RegexPatterns.KEY_VALUE_OR_FLOW.value, line)
            para_len = len(GeneralHelper.split_into_sentences(line))
            if para_len > 0:
                # count paragraph index, starting from 1 only if it consists at least a sentence
                paragraph_no = paragraph_no + 1
            for key_value_and_flow_pair in key_value_and_flow_pairs:
                if re.match(RegexPatterns.KEY_VALUE.value, key_value_and_flow_pair):
                    # returns tuple with key, value, measure, unit, log
                    metadata_set = cls.convert_bracketed_string_to_metadata(key_value_and_flow_pair)
                    # measure, unit, log could be empty
                    if metadata_set[4] != "":
                        log = log + "\n" + metadata_set[4]
                    if metadata_set[0] != "" and metadata_set[1] != "":
                        if re.search(RegexPatterns.COMMENT.value, metadata_set[0]):
                            key, comment = cls.process_internal_comment(metadata_set[0])
                            internal_comments.append(comment)
                        else:
                            key = metadata_set[0]
                        if re.search(RegexPatterns.COMMENT.value, metadata_set[1]):
                            val, comment = cls.process_internal_comment(metadata_set[1])
                            internal_comments.append(comment)
                        else:
                            val = metadata_set[1]
                        if re.search(RegexPatterns.COMMENT.value, metadata_set[2]):
                            measure, comment = cls.process_internal_comment(metadata_set[2])
                            internal_comments.append(comment)
                        else:
                            measure = metadata_set[2]
                        if re.search(RegexPatterns.COMMENT.value, metadata_set[3]):
                            unit, comment = cls.process_internal_comment(metadata_set[3])
                            internal_comments.append(comment)
                        else:
                            unit = metadata_set[3]
                        paragraph_key_pair = [paragraph_no, key]
                        if paragraph_key_pair in paragraph_key_pairs:
                            log = log + MiscAlertMsg.SIMILAR_PAR_KEY_FOUND.value.format(paragraph_key_pair) + "\n"
                        if cls.is_explicit_key(key):
                            key = TextCleaner.strip_colon(key)
                        metadata_pair = [paragraph_no, key, val, measure, unit]
                        paragraph_key_pairs.append(paragraph_key_pair)
                        metadata_pairs.append(metadata_pair)
                if re.match(RegexPatterns.FLOW.value, key_value_and_flow_pair):
                    flow_metadata, flow_log, is_flow_error = cls.extract_flow_type(paragraph_no,
                                                                                   key_value_and_flow_pair)
                    log = log + flow_log  # + "\n"
                    if is_flow_error:
                        break
                    metadata_pairs.extend(flow_metadata)
        print(log)
        return metadata_pairs, internal_comments, log

    # parse an opened document, first draft of sop
    @classmethod
    def convert_bracketed_string_to_metadata(cls, bracketed_str: str) -> Tuple[str, str, str, str, str]:
        """
        Extract lines to a tuple containing key, value, measure, and log.

        :param str bracketed_str: a string fragment with a single lister bracketing annotation.
        :returns: tuple (key, value, measure, unit, log)
            WHERE
            str key: the key portion of the string fragment,
            str value: the value portion of the string fragment,
            str measure: the measure portion of the string fragment,
            str unit: the unit portion of the string fragment,
            str log: log resulted from executing this and underlying functions.
        """
        log = ""
        bracketed_str_source = bracketed_str
        bracketed_str = bracketed_str[1:-1]
        split_metadata = re.split(r"\|", bracketed_str)
        if len(split_metadata) == 2:
            key = split_metadata[1]
            value = TextCleaner.strip_unwanted_mvu_colons(split_metadata[0])  # mvu: measure, value, unit
            measure = ""
            unit = ""
        elif len(split_metadata) == 3:
            value = TextCleaner.strip_unwanted_mvu_colons(split_metadata[0])
            unit = TextCleaner.strip_unwanted_mvu_colons(split_metadata[1])
            key = split_metadata[2]
            measure = ""
        elif len(split_metadata) == 4:
            measure = TextCleaner.strip_unwanted_mvu_colons(split_metadata[0])
            unit = TextCleaner.strip_unwanted_mvu_colons(split_metadata[1])
            key = split_metadata[3]
            value = TextCleaner.strip_unwanted_mvu_colons(split_metadata[2])
        elif len(split_metadata) == 1:
            key = ""
            value = ""
            measure = ""
            unit = ""
            print(MiscAlertMsg.SINGLE_PAIRED_BRACKET.value.format(bracketed_str_source))
            log = MiscAlertMsg.SINGLE_PAIRED_BRACKET.value.format(bracketed_str_source)
        else:
            log = MiscAlertMsg.INVALID_METADATA_SET_ELEMENT_NO.value.format(len(split_metadata),
                                                                            str(bracketed_str_source))
            raise SystemExit(log)
        key = key.strip()
        value = value.strip()
        measure = measure.strip()
        unit = unit.strip()
        return key, value, measure, unit, log

    @classmethod
    def conv_html_to_metadata(cls, html_content: str) -> Tuple[List, str]:
        """
        Turn html body content into an extended key-value pair
                [order, key, value, measure (if applicable), unit (if applicable)] or
                [-, section level, section name, '', ''].

        :param str html_content: body of the html content, extracted from eLabFTW API experiment.
        :return: tuple (multi_metadata_pair, log)
            WHERE
            list multi_metadata_pair is a list of a list with
                [order, key, value, measure (if applicable), unit (if applicable)] or
                [-, section level, section name, '', ''],
            str log is a string log returned from the respectively-executed functions.
        """
        # global log
        soup = BeautifulSoup(html_content.encode("utf-8"), "html.parser", from_encoding="found_encoding")
        soup.encoding = "utf-8"
        soup = TextCleaner.remove_table_tag(soup)

        clean_lines = TextCleaner.process_nbsp(soup)
        if clean_lines is not None:
            multi_metadata_pair, internal_comments, log = MetadataExtractor.parse_lines_to_metadata(clean_lines)
        else:
            multi_metadata_pair = None
            log = MiscAlertMsg.NO_HTML_LINE_CONTENT.value
            print(log)
        return multi_metadata_pair, log


class Validator:

    @classmethod
    def get_nonempty_body_tags(cls, exp: BeautifulSoup) -> List:
        """
        Clean up the source-html from empty-content html tags.

        :param bs4.soup exp: beautifulSoup4.soup experiment object.
        :return: list tagged_contents: list of non-empty html tags as well as new lines.
        """
        html_body = exp["body"]
        soup = BeautifulSoup(html_body.encode("utf-8"), "html.parser")
        soup.encoding = "utf-8"

        non_empty_soup = TextCleaner.remove_empty_tags(soup)
        tagged_contents = non_empty_soup.currentTag.tagStack[0].contents
        return tagged_contents

    @classmethod
    def check_bracket_num(cls, paragraph_no: int, text: str) -> Tuple[str, bool]:
        """
        Check if there is any bracketing error over the text line

        :param int paragraph_no: paragraph number for the referred line
        :param str text: string of the line
        :return: tuple (log, is_error)
            WHERE
            str log: error log (if any)
            bool is_error: flag if the checked line contains bracketing error
        """
        log = ""
        base_error_warning = "BRACKET ERROR: %s %s: %s"
        is_error = False
        if text.count("{") != text.count("}"):
            is_error = True
            log = base_error_warning % (BracketPairErrorMsg.IMPROPER_KEY_VALUE_BRACKET.value, str(paragraph_no), text)
        if text.count("<") != text.count(">"):
            is_error = True
            log = base_error_warning % (BracketPairErrorMsg.IMPROPER_FLOW_BRACKET.value, str(paragraph_no), text)
        if text.count("[") != text.count("]"):
            is_error = True
            log = base_error_warning % (BracketPairErrorMsg.IMPROPER_RANGE_BRACKET.value, str(paragraph_no), text)
        if text.count("(") != text.count(")"):
            is_error = True
            log = base_error_warning % (BracketPairErrorMsg.IMPROPER_COMMENT_BRACKET.value, str(paragraph_no), text)
        # print(log)
        return log, is_error

    # Used in several control flow validation functions.
    @classmethod
    def is_valid_comparative_operator(cls, operator: str) -> bool:
        """
        Check if the given operator is a valid comparative operator.

        :param str operator: The operator to check.
        :return: True if the operator is valid, False otherwise.
        """
        operators_list = ["e", "ne", "lt", "lte", "gt", "gte", "between"]
        if operator.lower() in operators_list:
            return True
        else:
            return False

    # Used in several control flow validation functions.
    @classmethod
    def is_valid_iteration_operator(cls, operator: str) -> bool:
        """
        Check if the given operator is a valid iteration operator.

        :param str operator: The operator to check.
        :return: True if the operator is valid, False otherwise.
        """
        operators_list = ["+", "-", "*", "/", "%"]
        if operator.lower() in operators_list:
            return True
        else:
            return False

    @classmethod
    def validate_while(cls, control_flow_split: List[str]) -> Tuple[str, bool]:
        """
        Validate the while command in the given list of strings.
        :param List[str] control_flow_split: List of strings containing the command and its arguments.
        :return: tuple (log, is_error)
            WHERE
            str log: error log (if any)
            bool is_error: flag if the checked line contains an error
        """
        log = ""
        is_error = False
        elements = len(control_flow_split)
        if elements == ArgNum.ARG_NUM_WHILE.value:
            if GeneralHelper.is_num(control_flow_split[1]):
                is_error = True
                log = log + MiscAlertMsg.ARGUMENT_MISMATCH.value.format(control_flow_split[1],
                                                                        control_flow_split) + "\n"
            if not cls.is_valid_comparative_operator(control_flow_split[2]):
                is_error = True
                log = log + MiscAlertMsg.UNRECOGNIZED_OPERATOR.value.format(control_flow_split[2],
                                                                            control_flow_split) + "\n"
        else:
            log = log + MiscAlertMsg.IMPROPER_ARGUMENT_NO.value.format(
                control_flow_split[0].upper(), ArgNum.ARG_NUM_WHILE.value, elements,
                control_flow_split) + "\n"
            is_error = True
        # note that the last value (comparison point is not yet checked as it can be a digit, binary or possibly
        # other things)
        return log, is_error

    # Used in process_foreach()
    @classmethod
    def validate_foreach(cls, control_flow_split: List[str]) -> Tuple[str, bool]:
        """
        Validate the foreach command in the given list of strings.

        :param List[str] control_flow_split: List of strings containing the command and its arguments.
        :return: tuple (log, is_error)
            WHERE
            str log: error log (if any)
            bool is_error: flag if the checked line contains an error
        """
        log = ""
        is_error = False
        elements = len(control_flow_split)
        if elements == ArgNum.ARG_NUM_FOREACH.value:
            if GeneralHelper.is_num(control_flow_split[1]):  # or
                # https://stackoverflow.com/questions/36330860/pythonically-check-if-a-variable-name-is-valid
                is_error = True
                log = log + MiscAlertMsg.ARGUMENT_MISMATCH.value.format(control_flow_split[1],
                                                                        control_flow_split) + "\n"
        else:
            log = log + MiscAlertMsg.IMPROPER_ARGUMENT_NO.value.format(
                control_flow_split[0].upper(), ArgNum.ARG_NUM_FOREACH.value, elements,
                control_flow_split) + "\n"
            is_error = True
        return log, is_error

    # Used in process_if().
    @classmethod
    def validate_if(cls, control_flow_split):
        """
        Validate the structure of an IF statement.

        This function checks the number of elements, the validity of the comparative operator,
        and the argument types in the provided IF statement.

        :param list control_flow_split: A list of elements in the IF statement.
        :return: tuple (log, is_error)
            WHERE
            str log: A log message.
            bool is_error: A flag indicating if there's an error.
        """

        log = ""
        is_error = False
        elements = len(control_flow_split)
        if elements == ArgNum.ARG_NUM_IF.value:
            if GeneralHelper.is_num(control_flow_split[1]):
                is_error = True
                log = log + MiscAlertMsg.ARGUMENT_MISMATCH.value.format(control_flow_split[1],
                                                                        control_flow_split) + "\n"
            if not cls.is_valid_comparative_operator(control_flow_split[2]):
                is_error = True
                log = log + MiscAlertMsg.UNRECOGNIZED_OPERATOR.value.format(control_flow_split[2],
                                                                            control_flow_split) + "\n"
        else:
            log = log + MiscAlertMsg.IMPROPER_ARGUMENT_NO.value.format(
                control_flow_split[0].upper(), ArgNum.ARG_NUM_IF.value, elements,
                control_flow_split) + "\n"
            is_error = True
        # note that the last value (comparison point) is not yet checked as it can be a digit, binary or possibly
        # other things
        return log, is_error

    # Used in process_elseif().
    # Validation functions for 'else-if', 'while', and 'if' have similar properties.
    # Hence, these functions can be integrated, but if there are changes for each of those, it may be challenging
    # to refactor.
    # For now, these validation functions are provided individually.
    @classmethod
    def validate_elseif(cls, control_flow_split):
        """
        Validate the structure of an ELSEIF statement.

        This class method checks the number of elements, the validity of the comparative operator,
        and the argument types in the provided ELSEIF statement.

        :param list control_flow_split: A list of elements in the ELSEIF statement.
        :return: tuple (log, is_error)
            WHERE
            str log: A log message.
            bool is_error: A flag indicating if there's an error.
        """
        log = ""
        is_error = False
        elements = len(control_flow_split)
        if elements == ArgNum.ARG_NUM_ELSEIF.value:
            if GeneralHelper.is_num(control_flow_split[1]):
                is_error = True
                log = log + MiscAlertMsg.ARGUMENT_MISMATCH.value.format(control_flow_split[1],
                                                                        control_flow_split) + "\n"
            if not cls.is_valid_comparative_operator(control_flow_split[2]):
                is_error = True
                log = log + MiscAlertMsg.UNRECOGNIZED_OPERATOR.value.format(control_flow_split[2],
                                                                            control_flow_split) + "\n"
        else:
            log = log + MiscAlertMsg.IMPROPER_ARGUMENT_NO.value.format(
                control_flow_split[0].upper(), ArgNum.ARG_NUM_ELSEIF.value, elements,
                control_flow_split) + "\n"
            is_error = True
        # note that the last value (comparison point is not yet checked as it can be a digit, binary or possibly
        # other things)
        return log, is_error

    # Used in else().
    @classmethod
    def validate_else(cls, control_flow_split: List[str]) -> Tuple[str, bool]:
        """
        Validate the else command in the given list of strings.

        :param List[str] control_flow_split: List of strings containing the command and its arguments.
        :return: tuple (log, is_error)
            WHERE
            str log: error log (if any)
            bool is_error: flag if the checked line contains an error
        """
        log = ""
        is_error = False
        elements = len(control_flow_split)
        if elements != ArgNum.ARG_NUM_ELSE.value:
            log = log + MiscAlertMsg.IMPROPER_ARGUMENT_NO.value.format(
                control_flow_split[0].upper(), ArgNum.ARG_NUM_ELSE.value, elements,
                control_flow_split) + "\n"
            is_error = True
        return log, is_error

    # Used in process_range().
    @classmethod
    def validate_range(cls, flow_range: str) -> Tuple[str, bool]:
        """
        Validate the range command in the given string.

        :param str flow_range: String containing the range command.
        :return: tuple (log, is_error)
            WHERE
            str log: error log (if any)
            bool is_error: flag if the checked line contains an error
        """
        is_error = False
        log = ""
        range_values = re.split("-", flow_range[1:-1])
        if len(range_values) == 2:
            if not (GeneralHelper.is_num(range_values[0]) and GeneralHelper.is_num(range_values[0])):
                is_error = True
                log = log + MiscAlertMsg.RANGE_NOT_NUMBERS.value.format(flow_range) + "\n"
        else:
            is_error = True
            log = log + MiscAlertMsg.RANGE_NOT_TWO_ARGUMENTS.value.format(flow_range) + "\n"
        return log, is_error

    # Used in process_for().
    @classmethod
    def validate_for(cls, control_flow_split: List[str]) -> Tuple[str, bool]:
        """
        Validate the 'for' iteration in the given list of strings.

        :param List[str] control_flow_split: List of strings containing the command and its arguments.
        :return: tuple (log, is_error)
            WHERE
            str log: error log (if any)
            bool is_error: flag if the checked line contains an error
        """
        log = ""
        is_error = False
        elements = len(control_flow_split)
        if elements == ArgNum.ARG_NUM_FOR.value:  # validating the number of arguments in FOR
            if GeneralHelper.is_num(control_flow_split[1]):  # in case 2nd argument is number, throw an error
                is_error = True
                log = log + MiscAlertMsg.ARGUMENT_MISMATCH.value.format(control_flow_split[1],
                                                                        control_flow_split) + "\n"
            range_error_log, is_range_error = cls.validate_range(control_flow_split[2])
            if is_range_error is True:  # check whether it is a valid range
                is_error = True
                log = log + range_error_log + "\n"
            if not cls.is_valid_iteration_operator(control_flow_split[3]):  # check whether it is a valid operator
                is_error = True
                log = log + MiscAlertMsg.INVALID_ITERATION_OPERATOR.value.format(control_flow_split[3],
                                                                                 control_flow_split) + "\n"
        else:  # if the number of parameters is invalid
            log = log + MiscAlertMsg.IMPROPER_ARGUMENT_NO.value.format(
                control_flow_split[0].upper(), ArgNum.ARG_NUM_FOR.value, elements,
                control_flow_split) + "\n"
            is_error = True
        return log, is_error

    # Used in process_iterate().
    @classmethod
    def validate_iterate(cls, control_flow_split: List[str]) -> Tuple[str, bool]:
        """
        Validate the iterate command in the given list of strings.

        :param List[str] control_flow_split: List of strings containing the command and its arguments.
        :return: tuple (log, is_error)
            WHERE
            str log: error log (if any)
            bool is_error: flag if the checked line contains an error
        """
        log = ""
        is_error = False
        elements = len(control_flow_split)
        if elements == ArgNum.ARG_NUM_ITERATE.value:
            if not cls.is_valid_iteration_operator(control_flow_split[1]):
                is_error = True
                log = log + MiscAlertMsg.INVALID_ITERATION_OPERATOR.value.format(control_flow_split[1],
                                                                                 control_flow_split) + "\n"
        else:  # if the number of parameters is invalid
            log = log + MiscAlertMsg.IMPROPER_ARGUMENT_NO.value.format(
                control_flow_split[0].upper(), ArgNum.ARG_NUM_ITERATE.value, elements,
                control_flow_split) + "\n"
            is_error = True
        return log, is_error

    # Used in process_section().
    @classmethod
    def validate_section(cls, control_flow_split: List[str]) -> Tuple[str, bool]:
        """
        Validate the section command in the given list of strings.

        :param List[str] control_flow_split: List of strings containing the command and its arguments.
        :return: tuple (log, is_error)
            WHERE
            str log: error log (if any)
            bool is_error: flag if the checked line contains an error
        """
        log = ""
        is_error = False
        elements = len(control_flow_split)
        if elements != ArgNum.ARG_NUM_SECTION.value:
            log = log + MiscAlertMsg.IMPROPER_ARGUMENT_NO.value.format(
                control_flow_split[0].upper(), ArgNum.ARG_NUM_SECTION.value,
                elements, control_flow_split) + "\n"
            is_error = True
        return log, is_error


class GeneralHelper:

    @classmethod
    def split_into_sentences(cls, content):
        """
        Split a line into proper sentences.

        :param str content: a line string that potentially consists of multiple sentences.
        :return: list sentences: list of split sentences, with regular/annotation bracket still intact.
        """
        # The code in this function is adapted from user:5133085's answer in SO:
        # https://stackoverflow.com/a/31505798/548451
        # (CC-BY-SA), see https://stackoverflow.com/help/licensing.
        latin_alphabets = "([A-Za-z])"
        openers = r"(Mr|Mrs|Ms|Dr|He\s|She\s|It\s|They\s|Their\s|Our\s|We\s|But\s|However\s|That\s|This\s|Wherever)"
        abbreviations = "([A-Z][.][A-Z][.](?:[A-Z][.])?)"
        prefixes = "(Mr|St|Mrs|Ms|Dr)[.]"
        sites = "[.](com|net|org|io|gov|de|eu)"
        suffixes = "(Inc|Ltd|Jr|Sr|Co)"
        content = " " + content + "  "
        content = content.replace("\n", " ")
        content = re.sub(prefixes, "\\1<prd>", content)
        content = re.sub(sites, "<prd>\\1", content)
        content = re.sub(r"\s" + latin_alphabets + "[.] ", " \\1<prd> ", content)
        content = re.sub(abbreviations + " " + openers, "\\1<stop> \\2", content)
        content = re.sub(latin_alphabets + "[.]" + latin_alphabets + "[.]" + latin_alphabets
                         + "[.]", "\\1<prd>\\2<prd>\\3<prd>", content)
        content = re.sub(latin_alphabets + "[.]" + latin_alphabets + "[.]", "\\1<prd>\\2<prd>", content)
        content = re.sub(" " + suffixes + "[.] " + openers, " \\1<stop> \\2", content)
        content = re.sub(" " + suffixes + "[.]", " \\1<prd>", content)
        content = re.sub(" " + latin_alphabets + "[.]", " \\1<prd>", content)
        if "”" in content:
            content = content.replace(".”", "”.")
        if "\"" in content:
            content = content.replace(".\"", "\".")
        if "!" in content:
            content = content.replace("!\"", "\"!")
        if "?" in content:
            content = content.replace("?\"", "\"?")
        content = content.replace(".", ".<stop>")
        content = content.replace("?", "?<stop>")
        content = content.replace("!", "!<stop>")
        content = content.replace("<prd>", ".")
        sentences = content.split("<stop>")
        sentences = sentences[:-1]
        sentences = [s.strip() for s in sentences]
        return sentences

    # Used in get_attachment_long_name()
    @classmethod
    def split_by_separators(cls, text: str, separators: List[str]) -> List[str]:
        """
        Split a given text using multiple separators.

        This function replaces all occurrences of the separators in the input
        string with the first separator in the list. Then, it splits the input
        string into a list of substrings based on the first separator. Finally,
        it returns the list of substrings after removing any leading or trailing
        whitespaces from each substring.

        :param str text: The input string to be split.
        :param list separators: A list of separators to be used for splitting the input string.
        :return: A list of substrings obtained by splitting the input string using the specified separators.
        :rtype: list
        """
        default_sep = separators[0]
        for sep in separators[1:]:
            text = text.replace(sep, default_sep)
        return [i.strip() for i in text.split(default_sep)]

    # Used in several control flow validation functions.
    @classmethod
    def is_num(cls, s: str) -> bool:
        """
        Check if the given string represents a number (integer or float).

        :param str s: The string to check.
        :return: True if the string represents a number, False otherwise.
        """
        if isinstance(s, int) or isinstance(s, float):
            return True
        else:
            s = s.replace(',', '', 1)
            if s[0] in ('-', '+'):
                return s[1:].isdigit()
            else:
                return s.isdigit()

    # helper function to print dataframe, used for development and debugging
    @classmethod
    def print_whole_df(cls, df: pd.DataFrame) -> None:
        """
        Print the entire DataFrame without truncation.
        """
        with pd.option_context('display.max_rows', None, 'display.max_columns',
                               None):  # more options can be specified also
            print(df)


class DocxHelper:

    # Used in write_tag_to_doc()
    @classmethod
    def get_span_attr_val(cls, c: Tag) -> Tuple[str, str]:
        """
        Get the attribute and value from the "style" attribute of a given Tag.

        This class method uses a regular expression to find the attribute and value
        in the "style" attribute of the provided Tag. It returns a tuple containing
        the attribute and value.

        :param bs4.element.Tag c: The Tag to extract the attribute and value from.
        :return: tuple (attr, val)
        WHERE
            str attr: The attribute.
            str val: The value.
        """
        found = re.findall(RegexPatterns.SPAN_ATTR_VAL.value, c.get("style"))
        attr, val = found[0]
        return attr, val

    # Used in write_tag_to_doc()
    @classmethod
    def get_section_title(cls, line: str) -> str:
        """
        Get the section title from a given line.

        This class method splits the given line into words and returns a string
        composed of all words except the first one. If the line contains only one word,
        an empty string is returned.

        :param str line: The line to extract the section title from.
        :return: The section title or an empty string.
        :rtype: str
        """
        words = line.split()
        if len(words) > 1:
            return ' '.join(words[1:])
        else:
            return ""

    @classmethod
    def process_reg_bracket(cls, line: str) -> Tuple[str, List[str]]:
        """
        Process strings with regular brackets (), which can be (_invisible comment_), (regular comment), or (DOI).
        This class method also maintains and updates the numerical index of DOIs found in the text entries.

        The string is returned to prepare for further docx content processing, in which the invisible comment will not
        be included, visible regular comment is still there but without brackets, and the DOI is provided with numerical
        index reference.

        :param str line: the comment string (with bracket) to be processed.
        :return: tuple (processed_line, references)
            WHERE
            str processed_line: processed_line is the processed string to be written as a part of docx content,
            list references: the list of available DOI references.
        """
        global ref_counter
        references = []
        # split based on the existence of brackets - including the captured bracket block in the result
        line_elements = re.split(RegexPatterns.COMMENT_W_CAPTURE_GROUP.value, line)
        processed_line = ""
        for element in line_elements:
            found_dois = re.findall(RegexPatterns.DOI.value, element)
            found_doi_length = len(found_dois)
            if re.search(RegexPatterns.COMMENT.value, element):
                # _invisible_ comment - strip all content (brackets, underscores, content
                if re.search(RegexPatterns.COMMENT_INVISIBLE.value, element):
                    processed_element = ""
                # visible comment - strip brackets and colons, keep the content
                elif re.search(RegexPatterns.COMMENT_VISIBLE.value, element):
                    element = re.sub(RegexPatterns.SEPARATOR_MARKUP.value, '', element)
                    visible_comments = re.split(RegexPatterns.COMMENT_VISIBLE.value, element)
                    concatenated_string = ""
                    for visible_comment in visible_comments:
                        concatenated_string = concatenated_string + visible_comment
                    processed_element = concatenated_string
                # comment that refers to DOI - strip all for now
                elif found_doi_length > 0:
                    processed_element = ""
                    for doi in found_dois:
                        ref_counter = ref_counter + 1
                        processed_element = processed_element + " [" + str(ref_counter) + "]"
                        references.append(doi)
                # otherwise, keep as is.
                else:
                    processed_element = element
            else:
                processed_element = element
            processed_line = processed_line + processed_element
        return processed_line, references

    # TODO: check why some invisible key elements passed the invisibility checks.
    @classmethod
    def write_tag_to_doc(cls, document: Document, tag_item: Tag) -> List[str]:
        """
        writes and format html tag content to a docx document.
        :param Document document: python-docx document instance.
        :param  bs4.element.Tag tag_item: tag to be processed and written to document.
        :return: list all_references
            WHERE
            list all_references: all references of DOIs contained in the document.
        """
        all_references = []
        p = document.add_paragraph()
        log = ""
        if isinstance(tag_item, Tag):
            section_toggle = False
            subsection_level = 0
            for subcontent in tag_item.contents:
                # strip_markup_and_explicit_keys()
                if isinstance(subcontent, Tag):
                    # print("ORIGINAL CONTENT OF SUBCONTENT.GETTEXT() WITHIN A TAG INSTANCE : "
                    # + subcontent.get_text())
                    line, references = TextCleaner.strip_markup_and_explicit_keys(subcontent.get_text())
                    # print("LINE FROM TAG INSTANCE: " + line)
                else:
                    if re.match(RegexPatterns.FORMULA.value, subcontent):
                        references = []
                        line = ""
                        formulas = re.findall(RegexPatterns.FORMULA.value, subcontent)
                        processed_subcontent = str(subcontent)
                        for formula in formulas:
                            stripped_formula = formula[1:-1]
                            processed_subcontent = processed_subcontent.replace(formula, '')
                            docx_formula, docx_formula_log = cls.latex_formula_to_docx(stripped_formula)
                            log = log + docx_formula_log
                            if docx_formula is not None:
                                p._element.append(docx_formula)
                                p.add_run(TextCleaner.remove_extra_spaces(processed_subcontent))
                    else:
                        line, references = TextCleaner.strip_markup_and_explicit_keys(subcontent.string)
                        # print("LINE FROM NON-TAG INSTANCE: " + line)

                if len(references) > 0:
                    all_references.extend(references)

                # check if the line is either goal, procedure, or result - but only limit that to one word
                if re.match(r'Goal:*|Procedure:*|Result:*', line, re.IGNORECASE) and len(line.split()) == 1:
                    document.add_heading(line, level=1)
                # check if the line is a section with the following characters
                elif re.match(RegexPatterns.SUBSECTION_W_EXTRAS.value, line, re.IGNORECASE):
                    section_title = cls.get_section_title(line)
                    subsection_level = line.count("sub")
                    if subsection_level == 0:
                        document.add_heading(section_title, level=2)
                    elif subsection_level == 1:
                        document.add_heading(section_title, level=3)
                    else:
                        document.add_heading(section_title, level=4)
                # check if it is a subscript text
                elif subcontent.name == "sub":
                    # sub_text = p.add_run(line + " ")
                    # print("SUB : " + line)
                    sub_text = p.add_run(TextCleaner.remove_extra_spaces(line))
                    sub_text.font.subscript = True
                # check if it is an italic-formatted text
                elif subcontent.name == "em":
                    # print("EM : " + line)
                    italic_text = p.add_run(TextCleaner.remove_extra_spaces(line))
                    # italic_text = p.add_run(" " + line + " ")
                    italic_text.font.italic = True
                # # check if it is a span tag
                elif subcontent.name == "span":
                    attr, val = cls.get_span_attr_val(subcontent)
                    # check if the line is a section without being followed by other characters, hence it needs
                    # the following chars from the next line as its section title. This case tends to happen when a
                    # section in a line is not covered within the same span tag hierarchy as its label.
                    if re.match(RegexPatterns.SUBSECTION.value, line, re.IGNORECASE):
                        section_toggle = True
                        subsection_level = line.count("sub")
                    elif section_toggle:
                        # do not use get_section_title() here as it will remove the first word of the line.
                        # the 'section' part has already been removed in this span section.
                        if subsection_level == 0:
                            document.add_heading(line, level=2)
                        elif subsection_level == 1:
                            document.add_heading(line, level=3)
                        else:
                            document.add_heading(line, level=4)
                        section_toggle = False
                    elif attr == "color":
                        # print("COLOR : " + line)
                        color_text = p.add_run(TextCleaner.remove_extra_spaces(line))
                        color_text.font.color.rgb = RGBColor.from_string(val[1:])
                    elif attr == "font-style" and attr == "italic":
                        # print("FONT STYLE AND ITALIC : " + line)
                        styled_text = p.add_run(TextCleaner.remove_extra_spaces(line))
                        styled_text.italic = True
                    else:
                        # print("NON SUBSECTION/HEADING/COLOR/FSTYLE : " + line)
                        p.add_run(TextCleaner.remove_extra_spaces(line))
                # check if it is a bold format
                elif subcontent.name == "strong":
                    # print("STRONG : " + line)
                    bold_text = p.add_run(TextCleaner.remove_extra_spaces(line))
                    bold_text.bold = True
                # check if it is a superscript format
                elif subcontent.name == "sup":
                    # print("SUP : " + line)
                    # super_text = p.add_run(line + " ")
                    super_text = p.add_run(TextCleaner.remove_extra_spaces(line))
                    super_text.font.superscript = True
                else:
                    # print("NON SUP/STRONG/SPAN/EM/SUB/SECTION/SUBSECTION : " + line)
                    p.add_run(TextCleaner.remove_extra_spaces(line))

        else:
            line, references = TextCleaner.strip_markup_and_explicit_keys(tag_item.string)
            if len(references) > 0:
                all_references.extend(references)
            # print("NON-TAG INSTANCE : " + line)
            p.add_run(TextCleaner.remove_extra_spaces(line))
            # print("*"*50)
        return all_references, log

    # Used in add_img_to_doc()
    @classmethod
    def get_text_width(cls, document: Document) -> float:
        """
        Return the text width (in mm) a given document's first section.

        This class method calculates the text width by subtracting the left and right margins
        from the page width of a given document's first section.
        The result is then divided by 36000 to convert the measurement to millimeters.

        :param docx.Document document: The document to calculate the text width for.
        :return: A floating point value representing the text width in millimeters.
        """
        section = document.sections[0]
        return (section.page_width - section.left_margin - section.right_margin) / 36000

    @classmethod
    def latex_formula_to_docx(cls, latex_formula: str) -> Tuple[str, str]:
        """
        Convert latex formula to docx formula.

        This function requires MML2OMML.XSL style sheet, which is normally shipped with Microsoft Office suite.
        The style sheet file should be placed in the same directory as config.json file. Please check LISTER's readme.

        :param str latex_formula: latex string to be converted to docx formula representation.
        :return: tuple (docx_formula, log)
            WHERE
            str docx_formula: formula represented in docx-string compatible that is going to be written to
                              the docx file.
            str log: error log (if any)
        """
        log = ""
        mathml = latex2mathml.converter.convert(latex_formula)
        tree = etree.fromstring(mathml)
        try:
            xslt = etree.parse('MML2OMML.XSL')  # please check whether the path on Mac is ok
            transform = etree.XSLT(xslt)
            new_dom = transform(tree)
            docx_formula = new_dom.getroot()
        except FileNotFoundError:
            docx_formula = None
            log = log + MiscAlertMsg.MISSING_MML2OMML.value
            print(log)
        return docx_formula, log

    @classmethod
    def add_table_to_doc(cls, doc: Document, content: Tag) -> None:
        """
        Add table content to docx instance.

        :param doc: python-docx instance of the modified document.
        :param bs4.Elements.Tag content: html table tag.
        """
        html_str_table = str(content.contents)[1:-1]
        dfs = pd.read_html(StringIO("<table>" + html_str_table + "</table>"))
        # read_html unfortunately does not retain styles/formatting, hence write your own html table parser if
        # formatting needs to be retained.
        df = dfs[0]
        # print_whole_df(df)
        t = doc.add_table(df.shape[0], df.shape[1], style="Light Grid Accent 3")

        # process table header, merge if it has similar value with the next cell and skip adding NaN value to the cell
        for h in range(df.shape[-1]):
            if not pd.isna(df.values[0, h]):
                t.cell(0, h).text = str(df.values[0, h])
            if h < (df.shape[-1] - 1):
                if df.values[0, h] == df.values[0, h + 1]:
                    t.cell(0, h).merge(t.cell(0, h + 1))
                    h = h + 1

        # process remaining table entries
        for i in range(1, df.shape[0]):
            for j in range(df.shape[-1]):
                if not pd.isna(df.values[i, j]):
                    t.cell(i, j).text = str(df.values[i, j])

    @classmethod
    def add_img_to_doc(cls, document: Document, real_name: str, path: str, image_hash: str) -> None:
        """
        Add image to the document file, based on upload experiment_id and image name when it was uploaded.

        :param Document document: the document object that is being modified.
        :param str real_name: real name of the image when it was uploaded to eLabFTW.
        :param str path: path to the image/attachment.
        :param str image_hash: hash of the image.
        """
        log = ""
        if real_name:
            if platform.system() == "Windows":
                img_saving_path = path + '\\attachments\\'
                sanitized_img_saving_path = sanitize_filepath(img_saving_path, platform="auto")
                docx_img_path = sanitized_img_saving_path + "\\" + image_hash + "_" + real_name
            else:
                img_saving_path = path + '/attachments/'
                sanitized_img_saving_path = sanitize_filepath(img_saving_path, platform="auto")
                docx_img_path = sanitized_img_saving_path + "/" + image_hash + "_" + real_name
            try:
                document.add_picture(docx_img_path, width=Mm(cls.get_text_width(document)))
            except Exception as e:
                log = log + MiscAlertMsg.INACCESSIBLE_ATTACHMENT.value.format(real_name, str(e))
            print(log)


# ---------------------------------------------- Text Cleaning Class ---------------------------------------------------
class TextCleaner:

    @classmethod
    def get_nonempty_body_tags(cls, exp: dict) -> List:
        """
        Clean up the source-html from empty-content html tags.

        :param dict exp: a dictionary of experiment object.
        :return: list tagged_contents: list of non-empty html tags as well as new lines.
        """
        html_body = exp.__dict__["_body"]
        soup = BeautifulSoup(html_body.encode("utf-8"), "html.parser")
        soup.encoding = "utf-8"
        non_empty_soup = cls.remove_empty_tags(soup)
        tagged_contents = non_empty_soup.currentTag.tagStack[0].contents
        return tagged_contents

    @classmethod
    # should probably be refactored to remove_nbsp for clarity
    def process_nbsp(cls, soup: BeautifulSoup) -> List[str]:
        """
        Remove non-break space (nbsp), and provide a 'clean' version of the lines.

        :param bs4.BeautifulSoup soup: soup object that is going to be cleaned up from nbsp.
        :return: list clean_lines lines without nbsp.
        """
        non_break_space = u'\xa0'
        text = soup.get_text().splitlines()
        # fetching the experiment paragraph, line by line
        lines = [x for x in text if x != '\xa0']  # Remove NBSP if it is on a single list element
        # Replace NBSP with space if it is inside the text
        line_no = 1
        clean_lines = []
        for line in lines:
            line = line.replace(non_break_space, ' ')
            clean_lines.append(line)
            line_no = line_no + 1
        return clean_lines

    @classmethod
    def strip_unwanted_mvu_colons(cls, word: str) -> str:
        """
        Remove the surrounding colon on word(s) within the annotation bracket if it belongs value/measure/unit category.

        :param str word: string with or without colons.
        :return: str word without colons.
        """
        if re.search(RegexPatterns.SURROUNDED_WITH_COLONS.value, word):
            print("Surrounding colons in the value/measure/unit {} is removed".format(word).encode("utf-8"))
            word = word[1:-1]  # if there are colons surrounding the word remains, remove it
        return word

    @classmethod
    def strip_markup_and_explicit_keys(cls, line: str) -> Tuple[str, List[str]]:
        """
        Strip keys that are marked as in visible (i.e., keys that are enclosed with colon) and extract any occurring
        pattern of DOI as reference, strip curly and angle brackets, reformat any annotation with regular bracket and
        fetch the DOI references, and strip unnecessary white spaces.

        :param bs4.element.NavigableString/str line: string to be inspected.
        :return: list of string containing DOI number.
        """
        stripped_from_explicit_keys = re.sub(RegexPatterns.SEPARATOR_AND_KEY.value, '', line)
        # print(stripped_from_explicit_keys)
        # strip curly and angle brackets
        stripped_from_markup = re.sub(RegexPatterns.BRACKET_MARKUPS.value, '', stripped_from_explicit_keys)
        # process based on the types within regular comment
        comments_based_strip, references = DocxHelper.process_reg_bracket(stripped_from_markup)
        # strip separator (pipe symbol)
        stripped_from_markup = re.sub(RegexPatterns.SEPARATOR_COLON_MARKUP.value, ' ', comments_based_strip)
        # strip unnecessary whitespaces
        stripped_from_trailing_spaces = re.sub(RegexPatterns.PRE_PERIOD_SPACES.value, '.', stripped_from_markup)
        stripped_from_trailing_spaces = re.sub(RegexPatterns.PRE_COMMA_SPACES.value, ',',
                                               stripped_from_trailing_spaces)
        # strip from trailing whitespaces
        # stripped_from_trailing_spaces = " ".join(stripped_from_trailing_spaces.split())
        return stripped_from_trailing_spaces, references

    # Used in parse_lines_to_metadata().
    @classmethod
    def strip_colon(cls, key: str) -> str:
        """
        Strip colon found on key string.

        This class method uses a regular expression to remove all colons from the provided string.

        :param str key: The string to remove colons from.
        :return: str stripped_key: The string with all colons removed.
        """
        stripped_key = re.sub(r'\:', '', key)
        return stripped_key

    @classmethod
    def remove_empty_tags(cls, soup: BeautifulSoup) -> BeautifulSoup:
        """
        Remove empty tags from a BeautifulSoup object.

        :param BeautifulSoup soup: The BeautifulSoup object to be processed.
        :return: BeautifulSoup soup: BeautifulSoup object with empty tags removed.
        """
        for x in soup.find_all():
            # if the text within a tag is empty, and tag name is not img/br/etc. and it is not img within p tag:
            if len(x.get_text(strip=True)) == 0 and x.name not in ['img', 'br', 'td', 'tr', 'table', 'h1', 'h2', 'h3',
                                                                   'h5', 'h6'] and len(x.select("p img")) == 0:
                x.extract()
        return soup

    @classmethod
    def remove_extra_spaces(cls, line: str) -> str:
        """
        Remove extra spaces from a given line.

        This class method uses a regular expression to replace all occurrences of multiple spaces
        in the provided string with a single space.

        :param str line: The string to remove extra spaces from.
        :return: (str) The string with all extra spaces removed.
        """
        return re.sub(' +', ' ', line)

    @classmethod
    def remove_table_tag(cls, soup: BeautifulSoup) -> BeautifulSoup:
        """
        Remove table tags and its content from the soup object.

        :param bs4.BeautifulSoup soup: bs4 soup object.
        :return: bs4.BeautifulSoup soup: BeautifulSoup object without a table tag, and its content.
        """
        for table in soup("table"):
            table.decompose()
        return soup


# ------------------------------------------------ Path Helper Class --------------------------------------------------
class PathHelper:
    @classmethod
    def derive_filename_from_experiment(cls, experiment: Union[elabapi_python.Experiment, Dict]) -> str:
        """
        Derive a file name from the experiment dictionary.

        This class method checks if the provided experiment is a dictionary.
        If it is, it retrieves the title from the dictionary.
        If it's not a dictionary, it retrieves the title from the experiment's attributes.
        The title is then converted to a slug which will be used as a file name.

        :param Union[elabapi_python.Experiment, Dict] experiment: The experiment to derive the file name from.
                                                           It Can be a dictionary or an object with a "_title"
                                                           attribute.
        :return: str filename_from_exp: The derived file name.
        """
        if isinstance(experiment, dict):
            experiment_title = experiment["title"]
        else:
            experiment_title = experiment.__dict__["_title"]
        filename_from_experiment = PathHelper.slugify(experiment_title)
        return filename_from_experiment

    @classmethod
    def get_default_output_path(cls, file_name: str) -> str:
        """
        Create an output path based on the home path (OS-dependent) and output file name.
        The home path is OS-dependent. On Windows/Linux, it is in the output directory as the script/executables.
        On macOS, it is in the users' Apps/lister/output/ directory.

        :param str file_name: file name for the output.
        :return: str output_path: the output path created from appending lister's output home directory and
                  output file name.
        """
        # enforce output path's base to be specific to ~/Apps/lister/ + output + filename
        if platform.system() == "Darwin":
            home = str(Path.home())
            output_path = home + "/Apps/lister/output/" + file_name
            print("OUTPUT PATH: %s" % output_path)
        # in windows and linux, use the executable's directory as a base to provide the outputs instead of home dir.
        else:
            current_path = pathlib.Path().resolve()
            if platform.system() == "Windows":
                # output_path = '\\\\?\\' + str(current_path) + "\output"
                output_path = str(current_path) + "\output"
            else:
                output_path = str(current_path) + "/output/"
        return output_path

    @classmethod
    def manage_output_path(cls, dir_name: str, file_name: str) -> str:
        """
        Get the output path according to the respective platform.

        If it is on macOS, return the dir_name (which has already been appended with output filename),
        on Windows/Linux, return the dir_name + output file_name.

        :param str dir_name: the home directory name for the output.
        :param str file_name: the output name.
        :return: str output_path is the output directory created from appending the home path and output path.
        """
        # on macOS, enforce output path's base to be specific to ~/Apps/lister/ + output + filename
        if platform.system() == "Darwin":
            output_path = dir_name + file_name + "/"
        # in windows and linux, use the executable's directory as a base to provide the outputs instead of home dir‚
        else:

            if platform.system() == "Windows":
                # Prepend the '\\?\' prefix to allow long file paths on Windows
                base_path = dir_name + "\\" + file_name + "\\"
                output_path = "\\\\?\\" + base_path
            else:
                base_path = dir_name + "/" + file_name + "/"
                output_path = base_path

        return output_path

    @classmethod
    def check_and_create_path(cls, path: str) -> None:
        """
        Check if the given path exists, and create the directory if it doesn't.

        :param path: The path to check and create if necessary.
        """
        if not os.path.isdir(path):
            print("Output path %s is not available, creating the path directory..." % path)
            os.makedirs(path)

    @classmethod
    def manage_input_path(cls) -> str:
        """
        Enforce reading input from a specific directory on macOS (on macOS, LISTER cannot get the input directly
        from the executable file's directory).

        :return: str input_path is the input directory for macOS.
        """
        input_path = ""
        if platform.system() == "Darwin":  # enforce the input path to be specific to ~/Apps/lister/
            home = str(Path.home())
            input_path = home + "/Apps/lister/"
        return input_path

    #    Taken from https://github.com/django/django/blob/master/django/utils/text.py
    #    Convert to ASCII if 'allow_unicode' is False. Convert spaces or repeated
    #    dashes to single dashes. Remove characters that aren't alphanumerics,
    #    underscores, or hyphens. Convert to lowercase. Also strip leading and
    #   trailing whitespace, dashes, and underscores.
    @classmethod
    def slugify(cls, value: Union[str, Any], allow_unicode: bool = False) -> str:

        """
        Convert a string into a URL-friendly slug.

        :param value: The input string to be converted.
        :param allow_unicode: Whether to allow Unicode characters in the slug.
        :return: The URL-friendly slug.
        """
        value = str(value)
        if allow_unicode:
            value = unicodedata.normalize('NFKC', value)
        else:
            value = unicodedata.normalize('NFKD', value).encode('ascii', 'ignore').decode('ascii')
        value = re.sub(r'[^\w\s-]', '', value.lower())
        return re.sub(r'[-\s]+', '-', value).strip('-_')


# ------------------------------------------------ MAIN FUNCTION ------------------------------------------------------
ref_counter = 0


def main():
    global output_filename
    global output_path, base_output_path
    global token, exp_no, endpoint

    # suppress the redundant window pop up on macOS as a workaround, see
    # https://stackoverflow.com/questions/72636873/app-package-built-by-pyinstaller-fails-after-it-uses-tqdm
    if platform.system() == "Darwin":
        warnings.filterwarnings("ignore")
        freeze_support()

    gui_helper = GUIHelper()

    args = gui_helper.parse_gooey_args()
    base_output_path = args.base_output_dir
    api_v2_endpoint = ApiAccess.get_api_v2_endpoint(args.endpoint)
    api_v2_client = ApiAccess.create_api_v2_client(api_v2_endpoint, args.token)

    if args.command == 'parse_resource':
        item_api_response, resource_log = ApiAccess.get_resource_item(api_v2_client, args.resource_item_no)
        cat = item_api_response.__dict__["_category_title"]
        title = item_api_response.__dict__["_title"]
        if args.experiment_id:
            output_filename = PathHelper.slugify(cat) + "_" + str(args.resource_item_no)
        elif args.title:
            output_filename = PathHelper.slugify(cat) + "_" + PathHelper.slugify(title)
    elif args.command == 'parse_experiment':
        if args.experiment_id:
            output_filename = PathHelper.slugify("experiment") + "_" + str(args.exp_no)
        elif args.title:
            title = ApiAccess.get_exp_title(api_v2_client, args.exp_no)
            output_filename = PathHelper.slugify("experiment") + "_" + PathHelper.slugify(title)
    print("The output is written to %s directory" % output_filename)

    output_path = PathHelper.manage_output_path(args.base_output_dir, output_filename)
    PathHelper.check_and_create_path(output_path)

    print("base_output_dir: ", base_output_path)
    print("output_filename: ", output_filename)
    print("output_path: ", output_path)

    if args.command == 'parse_experiment':
        print("Processing an experiment...")
        MetadataExtractor.process_experiment(api_v2_client, args.exp_no, output_path)
    elif args.command == 'parse_resource':
        print("Processing a resource...")
        MetadataExtractor.process_ref_resource_item(api_v2_client, item_api_response)


if __name__ == "__main__":
    main()
