import elabapy
import ssl
from bs4 import BeautifulSoup, Tag
from enum import Enum
import re
from docx import Document


ref_counter = 0

class Regex_patterns(Enum):
    EXPLICIT_KEY = r':.+?:' # catch explicit key which indicated within ":" sign
    KV_OR_FLOW = r'\{.+?\}|<.+?>'  # find any occurrences of either KV or control flow
    KV = r'\{.+?\}'  # find any occurrences of KV
    FLOW = r'<.+?>'  # find any occurrences of control flows
    DOI = r"\b(10[.][0-9]{4,}(?:[.][0-9]+)*/(?:(?![\"&\'<>])\S)+)\b" # catch DOI
    COMMENT = "\(.+?\)"  # define regex for parsing comment
    COMMENT_W_CAPTURE_GROUP = "(\(.+?\))"
    COMMENT_VISIBLE = "\(:.+?:\)"
    COMMENT_INVISIBLE = "\(_.+?_\)"
    SEPARATOR_AND_KEY = r"\|(\s*\w\s*\.*)+\}" # catch the end part of KV pairs (the key, tolerating trailing spaces)
    BRACKET_MARKUPS = r"([{}<>])" # catch KV/section bracket annotations
    SEPARATOR_COLON_MARKUP = r"([|:])" # catch separator annotation
    PRE_PERIOD_SPACES = '\s+\.'
    PRE_COMMA_SPACES = '\s+,'
    SUBSECTION = '(sub)*section'
    SUBSECTION_W_EXTRAS = r'(sub)*section.+'


token = "805e37315e9dd3697543f1bad11c750123d4071b5ced3195b617a41583bc456a97a8172721f44dd7622a"
url_endpoint = "http://elabftw.hhu.de/api/v1/"
exp_id = 19787


def strip_markup_and_explicit_keys(line):
    # strip keys that are not marked visible (keys that are not enclosed with colon)
    stripped_from_explicit_keys = re.sub(Regex_patterns.SEPARATOR_AND_KEY.value, '', line)
    # strip curly and angle brackets
    stripped_from_markup = re.sub(Regex_patterns.BRACKET_MARKUPS.value, '', stripped_from_explicit_keys)
    # process based on the types within regular comment
    comments_based_strip = process_reg_bracket(stripped_from_markup)
    # strip separator (pipe symbol)
    stripped_from_markup = re.sub(Regex_patterns.SEPARATOR_COLON_MARKUP.value, ' ', comments_based_strip)
    # strip unnecessary whitespaces
    stripped_from_trailing_spaces = re.sub(Regex_patterns.PRE_PERIOD_SPACES.value, '.', stripped_from_markup)
    stripped_from_trailing_spaces = re.sub(Regex_patterns.PRE_COMMA_SPACES.value, ',', stripped_from_trailing_spaces)
    stripped_from_trailing_spaces = " ".join(stripped_from_trailing_spaces.split()) # strip from trailing whitespaces
    return stripped_from_trailing_spaces

def process_reg_bracket(line):
    global ref_counter
    # split based on the existence of brackets - including the captured bracket block in the result
    line_elements = re.split(Regex_patterns.COMMENT_W_CAPTURE_GROUP.value, line)
    processed_elements = []
    processed_line = ""
    for element in line_elements:
        # print(element)
        if re.search(Regex_patterns.COMMENT.value, element):
            # _invisible_ comment - strip all content (brackets, underscores, content
            if re.search(Regex_patterns.COMMENT_INVISIBLE.value, element):
                processed_element = ""
            # visible comment - strip brackets and colons, keep the content
            elif re.search(Regex_patterns.COMMENT_VISIBLE.value, element):
                processed_element = element[2:-2]
            # comment that refer to DOI - strip all for now
            elif re.search(Regex_patterns.DOI.value, element[1:-1]):
                ref_counter = ref_counter + 1
                processed_element = " [" + str(ref_counter) + "]"
            # otherwise, keep as is.
            else:
                processed_element = element
        else:
            processed_element = element
        processed_line = processed_line + processed_element
    return processed_line


def serialize_to_docx(narrative_lines, references):
    # print(type(narrative_lines))
    # here line has been stripped out of its tags, e.g.,:
    # ['Goal', 'Cooking a simple spaghetti con le acciughe that can be reproduced by beginner level cooks who wish to
    # cook pescatarian. This recipe uses spaghetti as the main ingredient.', 'Procedure', 'Section Initial Process',
    # '500 grams of spaghetti is cooked by boiling, using salted water as the boiling medium.', 'Section Sauce',
    # 'Subsection Sauce', 'Heating with a high heat level is done on 0.33 cups of extra virgin olive oil.', ...
    document = Document()
    reference_switch = False
    intext_reference_list = []
    for line in narrative_lines:
        # check if the line is either goal, procedure, result, or reference
        if re.match(r'Goal:*|Procedure:*|Result:*', line, re.IGNORECASE):
            document.add_heading(line, level=1)
            reference_switch = False
        # check if the line is a section
        # elif re.match(r'Section.+', line, re.IGNORECASE):
        elif re.match(Regex_patterns.SUBSECTION_W_EXTRAS.value, line, re.IGNORECASE):
            subsection_level = line.count("sub")
            line = re.sub(Regex_patterns.SUBSECTION_W_EXTRAS.value, '', line)
            if subsection_level == 0:
                document.add_heading(line.strip(), level=2)
            elif subsection_level == 1:
                document.add_heading(line.strip(), level=3)
            else:
                document.add_heading(line.strip(), level=4)
            reference_switch = False
        # check if the line is a reference
        elif re.match(r'References:*|Reference:*', line, re.IGNORECASE):
            # document.add_heading(line, level=1)
            reference_switch = True
        else:
            line = re.sub('\s{2,}', ' ', line) # replace superfluous whitespaces in preceding text with a single space
            line = re.sub(r'\s([?.!"](?:\s|$))', r'\1', line)
            if reference_switch == False:
                document.add_paragraph(line)
            else:
                intext_reference_list.append(line)
    # add reference list
    if reference_switch == True:
        document.add_heading("Reference", level=1)
        for intext_reference in intext_reference_list:
            document.add_paragraph(intext_reference, style="List Number")
    if len(references)>0:
        if reference_switch == False:
            document.add_heading("Reference", level=1)
        for reference in references:
            document.add_paragraph(reference, style='List Number')
    document.save(output_file_prefix + '.docx')

ssl._create_default_https_context = ssl._create_unverified_context
manager = elabapy.Manager(endpoint=url_endpoint, token=token, verify=False)
exp = manager.get_experiment(exp_id)

html_body = exp["body"]
soup = BeautifulSoup(html_body, "html.parser")

for x in soup.find_all():
    if len(x.get_text(strip=True)) == 0 and x.name not in ['br', 'img']:
        x.extract()



from docx import Document
document = Document()
#print(soup)
n = 0
body_tag_contents = soup.currentTag.tagStack[0].contents
references = []
for content in body_tag_contents:
    # if tag is p
    # use the strip function
    # if tag is table
    # process table
    if isinstance(content, Tag):
    # if type(content) is
       # print("ITS A TAG")
        if content.name == "p":
            print("ITS A PARAGRAPH")

            # print(type(content.contents))
            print(len(content.contents))
            # print(content.contents[0])
            print(type(str(content.string)))
            # print(str(content.string))
            line = strip_markup_and_explicit_keys(str(content.string))

            internal_comments = []
            # Extract overall comments, including those within KV pairs
            overall_comments = re.findall(Regex_patterns.COMMENT.value, line)

            print(line)
            # cleaned_line = strip_markup_and_explicit_keys(content.)
            if re.match(r'Goal:*|Procedure:*|Result:*', line, re.IGNORECASE):
                document.add_heading(line, level=1)
                reference_switch = False
            # check if the line is a section
            # elif re.match(r'Section.+', line, re.IGNORECASE):
            elif re.match(Regex_patterns.SUBSECTION_W_EXTRAS.value, line, re.IGNORECASE):
                subsection_level = line.count("sub")
                line = re.sub(Regex_patterns.SUBSECTION_W_EXTRAS.value, '', line)
                if subsection_level == 0:
                    document.add_heading(line.strip(), level=2)
                elif subsection_level == 1:
                    document.add_heading(line.strip(), level=3)
                else:
                    document.add_heading(line.strip(), level=4)
                reference_switch = False
            # check if the line is a reference
            elif re.match(r'References:*|Reference:*', line, re.IGNORECASE):
                # document.add_heading(line, level=1)
                reference_switch = True
            else:
                line = re.sub('\s{2,}', ' ',
                              line)  # replace superfluous whitespaces in preceding text with a single space
                line = re.sub(r'\s([?.!"](?:\s|$))', r'\1', line)
                if reference_switch == False:
                    document.add_paragraph(line)
                else:
                    intext_reference_list.append(line)
        elif content.name == "table":
            print("ITS A TABLE")
            # print(content.contents)
        n = n+1
    # add reference list
    if reference_switch == True:
        document.add_heading("Reference", level=1)
        for intext_reference in intext_reference_list:
            document.add_paragraph(intext_reference, style="List Number")
    if len(references)>0:
        if reference_switch == False:
            document.add_heading("Reference", level=1)
        for reference in references:
            document.add_paragraph(reference, style='List Number')
    document.save('output.docx')
        #print(type(content))
        #print(content)
        #print(n)